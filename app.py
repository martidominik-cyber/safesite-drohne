import streamlit as st
import google.generativeai as genai
import cv2
import tempfile
import os
import json
from fpdf import FPDF
import time
from datetime import date
from PIL import Image, ImageDraw, ImageFont
import urllib.parse
import uuid 

# Word-Modul sicher laden
try:
    from docx import Document
    from docx.shared import Inches
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# ==========================================
# 0. KONFIGURATION
# ==========================================
st.set_page_config(page_title="SafeSite Drohne", page_icon="logo.jpg", layout="wide", initial_sidebar_state="auto")

# ----------------------------------------------------
# 🔴 HIER DEINEN GITHUB-NAMEN EINTRAGEN!
LOGO_URL_GITHUB = "https://raw.githubusercontent.com/DEIN_BENUTZERNAME/safesite-drohne/main/logo.jpg?v=1"
# ----------------------------------------------------

# STYLE
st.markdown(f"""
<style>
    .stAppDeployButton {{display: none;}}
    footer {{visibility: hidden;}}
    [data-testid="stSidebarCollapsedControl"] {{color: #FF6600 !important;}}
    h1, h2, h3 {{ color: #FF6600 !important; }}
</style>
<link rel="apple-touch-icon" href="{LOGO_URL_GITHUB}">
""", unsafe_allow_html=True)

# DATENBANK
USER_DB_FILE = "users.json"
CUSTOMERS_DB_FILE = "customers.json"
GEFAHRSOFF_DB_FILE = "gefahrstoffe.json"

def load_users():
    if not os.path.exists(USER_DB_FILE):
        with open(USER_DB_FILE, "w") as f: json.dump({"admin": "1234"}, f)
    with open(USER_DB_FILE, "r") as f: return json.load(f)

def save_users(users):
    with open(USER_DB_FILE, "w") as f: json.dump(users, f, indent=2)

def load_customers():
    if not os.path.exists(CUSTOMERS_DB_FILE):
        with open(CUSTOMERS_DB_FILE, "w") as f: json.dump({}, f)
    with open(CUSTOMERS_DB_FILE, "r") as f: return json.load(f)

def save_customers(customers):
    with open(CUSTOMERS_DB_FILE, "w") as f: json.dump(customers, f, indent=2)

def load_gefahrstoffe():
    if not os.path.exists(GEFAHRSOFF_DB_FILE):
        with open(GEFAHRSOFF_DB_FILE, "w") as f: json.dump({}, f)
    with open(GEFAHRSOFF_DB_FILE, "r") as f: 
        gefahrstoffe = json.load(f)
        # Beim ersten Laden: Standard-Gefahrstoffe hinzufügen, wenn Liste leer ist
        if not gefahrstoffe:
            gefahrstoffe = init_standard_gefahrstoffe()
            save_gefahrstoffe(gefahrstoffe)
        return gefahrstoffe

def init_standard_gefahrstoffe():
    """Initialisiert die Standard-Gefahrstoffe basierend auf dem Muster-Kataster"""
    standard_gefahrstoffe = {}
    
    # 1. Zementhaltige Produkte
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "Zementhaltige Produkte",
        "handelsbezeichnung": "Beton, Mörtel, Fugenmassen",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "Zementhaltige Produkte",
        "cas_nummer": "",
        "lagerort": "Baustelle / Lager",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS05, GHS07",
        "gefahrenbeschreibung": "Ätzend / Reizend. Verursacht schwere Augenschäden. Hautreizungen (Maurerkrätze). Staub reizt Atemwege.",
        "schutzmassnahmen": "Handschuhe (Nitril/Butyl), Schutzbrille, lange Kleidung. Bei Staubentwicklung: Maske FFP2.",
        "verwendung": "Bauarbeiten mit Beton, Mörtel und Fugenmassen",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    # 2. Lösungsmittelhaltige Farben/Lacke/Kleber
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "Lösungsmittelhaltige Farben/Lacke/Kleber",
        "handelsbezeichnung": "Verdünner, Kunstharzlacke",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "Lösungsmittelhaltige Farben/Lacke/Kleber",
        "cas_nummer": "",
        "lagerort": "Giftschrank / Lager",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS02, GHS08, GHS07",
        "gefahrenbeschreibung": "Entzündbar / Gesundheitsschädlich. Dämpfe können Benommenheit verursachen. Kann Organe schädigen (Nervensystem).",
        "schutzmassnahmen": "Gute Lüftung sicherstellen. Zündquellen fernhalten. Atemschutzmaske (Filter Typ A - braun).",
        "verwendung": "Lackieren, Kleben mit lösungsmittelhaltigen Produkten",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "Prüfung: Wasserverdünnbare Alternativen verwenden, wenn möglich",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    # 3. Epoxidharze (2-Komponenten)
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "Epoxidharze (2-Komponenten)",
        "handelsbezeichnung": "Bodenbeschichtung, Injektionsmörtel",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "Epoxidharze (2-Komponenten)",
        "cas_nummer": "",
        "lagerort": "Giftschrank / Lager",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS09, GHS05, GHS07",
        "gefahrenbeschreibung": "Sensibilisierend / Gewässergefährdend. Starke allergische Hautreaktionen möglich. Giftig für Wasserorganismen.",
        "schutzmassnahmen": "Hautkontakt strikt vermeiden! Lange Ärmel, spezielle Einweghandschuhe (dicker Nitril), Schutzbrille.",
        "verwendung": "Bodenbeschichtungen, Injektionsarbeiten",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    # 4. PU-Produkte (Isocyanate)
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "PU-Produkte (Isocyanate)",
        "handelsbezeichnung": "Bauschaum, Montageschaum, PU-Kleber",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "PU-Produkte (Isocyanate)",
        "cas_nummer": "",
        "lagerort": "Giftschrank / Lager",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS08, GHS07, GHS02",
        "gefahrenbeschreibung": "Krebserzeugungsverdacht / Atemwegssensibilisierend. Kann bei Einatmen Allergien/Asthma auslösen. Extrem entzündbares Aerosol.",
        "schutzmassnahmen": "Gute Lüftung. Schutzhandschuhe. Bei schlechter Lüftung Atemschutz notwendig.",
        "verwendung": "Montagearbeiten, Dichtungsarbeiten, Kleben",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    # 5. Kraftstoffe & Schmiermittel
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "Kraftstoffe & Schmiermittel",
        "handelsbezeichnung": "Diesel, Benzin, Schalöl",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "Kraftstoffe & Schmiermittel",
        "cas_nummer": "",
        "lagerort": "Tankstelle / Auffangwanne",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS02, GHS08, GHS09",
        "gefahrenbeschreibung": "Entzündbar / Aspirationsgefahr. Kann tödlich sein bei Verschlucken/Eindringen in Atemwege. Umweltgefährlich.",
        "schutzmassnahmen": "Auffangwannen nutzen (Gewässerschutz). Feuerlöscher bereitstellen. Nicht rauchen.",
        "verwendung": "Betankung von Maschinen, Schalung",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    # 6. Reinigungsmittel (Sauer)
    standard_gefahrstoffe[str(uuid.uuid4())] = {
        "name": "Reinigungsmittel (Sauer)",
        "handelsbezeichnung": "Zementschleierentferner, Sanitärreiniger",
        "hersteller": "Verschiedene Hersteller",
        "kategorie": "Reinigungsmittel (Sauer)",
        "cas_nummer": "",
        "lagerort": "Giftschrank / Lager",
        "menge": "Variabel",
        "sdb_datum": "",
        "ghs_symbole": "GHS05",
        "gefahrenbeschreibung": "Korrosiv / Ätzend. Verursacht schwere Verätzungen der Haut und Augenschäden.",
        "schutzmassnahmen": "Schutzbrille (Korbbrille) zwingend. Säurebeständige Handschuhe.",
        "verwendung": "Reinigung von Zementschleiern, Sanitärreinigung",
        "betriebsanweisung_vorhanden": "Ja",
        "substitution": "",
        "sdb_link": "",
        "sdb_datei": "",
        "erstellt_am": date.today().strftime('%d.%m.%Y')
    }
    
    return standard_gefahrstoffe

def save_gefahrstoffe(gefahrstoffe):
    with open(GEFAHRSOFF_DB_FILE, "w") as f: json.dump(gefahrstoffe, f, indent=2)

def is_admin():
    return st.session_state.logged_in and st.session_state.username == "admin"

def get_customer_by_email(email):
    """Findet einen Kunden anhand seiner Email-Adresse"""
    customers = load_customers()
    for kunde_id, kunde_data in customers.items():
        if kunde_data.get('email') == email:
            return kunde_id, kunde_data
    return None, None

def get_customer_by_username_or_email(username_or_email):
    """Findet einen Kunden anhand Email oder Benutzername"""
    customers = load_customers()
    for kunde_id, kunde_data in customers.items():
        # Prüfe Email
        if kunde_data.get('email') == username_or_email:
            return kunde_id, kunde_data
        # Prüfe Benutzername
        if kunde_data.get('username') and kunde_data.get('username') == username_or_email:
            return kunde_id, kunde_data
    return None, None

def get_customer_email_from_login(username_or_email):
    """Gibt die Email eines Kunden zurück basierend auf Login-Username oder Email"""
    kunde_id, kunde_data = get_customer_by_username_or_email(username_or_email)
    if kunde_data:
        return kunde_data.get('email', username_or_email)
    return username_or_email

def get_customer_credits(email_or_username):
    """Gibt die Credits eines Kunden zurück (0 falls nicht gefunden) - akzeptiert Email oder Username"""
    kunde_id, kunde_data = get_customer_by_username_or_email(email_or_username)
    if kunde_data:
        return int(kunde_data.get('credits', 0))
    return 0

def deduct_credit(email_or_username):
    """Zieht 1 Credit vom Kunden ab und speichert - akzeptiert Email oder Username"""
    customers = load_customers()
    kunde_id, kunde_data = get_customer_by_username_or_email(email_or_username)
    if kunde_id and kunde_data:
        current_credits = int(kunde_data.get('credits', 0))
        if current_credits > 0:
            customers[kunde_id]['credits'] = current_credits - 1
            save_customers(customers)
            return True
    return False

def update_customer_credits(kunde_id, credits):
    """Aktualisiert die Credits eines Kunden"""
    customers = load_customers()
    if kunde_id in customers:
        customers[kunde_id]['credits'] = int(credits)
        save_customers(customers)
        return True
    return False

# API KEY CHECK
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except:
    st.error("⚠️ API Key fehlt in den Secrets!")
    st.stop()

# DATEIEN
LOGO_FILE = "logo.jpg"
TITELBILD_FILE = "titelbild.png"

# ==========================================
# 1. FUNKTIONEN
# ==========================================
def clean_json(text):
    text = text.strip()
    first_bracket = text.find('[')
    last_bracket = text.rfind(']')
    if first_bracket != -1 and last_bracket != -1:
        text = text[first_bracket:last_bracket+1]
    return text

def extract_frame(video_path, timestamp):
    try:
        cap = cv2.VideoCapture(video_path)
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
        ret, frame = cap.read(); cap.release()
        if ret: return cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    except: return None

def convert_image_if_needed(img_path):
    """Konvertiert Bilder in ein Format, das von PIL verarbeitet werden kann"""
    try:
        # Prüfe ob es eine HEIC/HEIF Datei ist
        if img_path.lower().endswith(('.heic', '.heif')):
            # Versuche zuerst mit PIL (falls pillow-heif installiert ist)
            try:
                img = Image.open(img_path)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                new_path = img_path.rsplit('.', 1)[0] + '.jpg'
                img.save(new_path, 'JPEG', quality=95)
                # Alte Datei löschen
                if os.path.exists(img_path):
                    try: os.remove(img_path)
                    except: pass
                return new_path
            except:
                # PIL kann HEIC nicht öffnen, versuche mit OpenCV
                pass
            
            # Versuche mit OpenCV (kann manchmal HEIC lesen, wenn entsprechende Codecs vorhanden sind)
            try:
                img_array = cv2.imread(img_path, cv2.IMREAD_COLOR)
                if img_array is not None and img_array.size > 0:
                    new_path = img_path.rsplit('.', 1)[0] + '.jpg'
                    cv2.imwrite(new_path, img_array, [cv2.IMWRITE_JPEG_QUALITY, 95])
                    if os.path.exists(img_path):
                        try: os.remove(img_path)
                        except: pass
                    return new_path
            except:
                pass
            
            # Falls beides fehlschlägt, gib Warnung aus aber behalte Originaldatei
            # (möglicherweise unterstützt der Browser die Konvertierung beim Upload)
            return img_path
        
        # Für andere Formate, versuche einfach zu öffnen
        try:
            img = Image.open(img_path)
            # Stelle sicher, dass es RGB ist
            if img.mode != 'RGB' and img.mode not in ['RGBA', 'P']:
                # Konvertiere problematische Formate
                if img.mode in ['RGBA', 'P']:
                    # Erstelle weissen Hintergrund für transparente Bilder
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        rgb_img.paste(img, mask=img.split()[3])
                    else:
                        rgb_img.paste(img)
                    new_path = img_path.rsplit('.', 1)[0] + '_rgb.jpg'
                    rgb_img.save(new_path, 'JPEG', quality=95)
                    if img_path != new_path and os.path.exists(img_path):
                        try: os.remove(img_path)
                        except: pass
                    return new_path
            return img_path
        except:
            return img_path
            
    except Exception as e:
        # Bei jedem Fehler, gib die Originaldatei zurück
        return img_path

# --- PDF GENERATOR ---
class PDF(FPDF):
    def header(self):
        # Logo oben RECHTS platzieren
        if os.path.exists(LOGO_FILE):
            try: self.image(LOGO_FILE, 160, 8, 40)
            except: pass
        self.ln(5)

def make_safe_text(text):
    """Entfernt Emojis für das PDF, damit es nicht abstürzt"""
    if text is None: return ""
    return text.encode('latin-1', 'ignore').decode('latin-1')

def create_pdf(data, m_type, m_files, projekt, inspektor, status):
    pdf = PDF()
    pdf.add_page()
    
    # --- HEADER BEREICH ---
    pdf.set_font("Arial", 'B', 20)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, "SICHERHEITS-INSPEKTION (DROHNE)", ln=True)
    pdf.ln(8)
    
    # Metadaten
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(35, 8, "Projekt:", ln=0)
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, make_safe_text(projekt), ln=True)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(35, 8, "Datum:", ln=0)
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, date.today().strftime('%d.%m.%Y') + f" | {time.strftime('%H:%M')} Uhr", ln=True)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(35, 8, "Inspektor:", ln=0)
    pdf.set_font("Arial", '', 11)
    pdf.cell(0, 8, make_safe_text(f"{inspektor} (SafeSite Drohne)"), ln=True)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(35, 8, "Status:", ln=0)
    pdf.set_font("Arial", '', 11)
    
    if "Massnahmen" in status: pdf.set_text_color(255, 153, 51)
    else: pdf.set_text_color(0, 153, 0)
    pdf.cell(0, 8, make_safe_text(status), ln=True)
    pdf.set_text_color(0, 0, 0) 
    
    pdf.ln(10)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "1. ZUSAMMENFASSUNG / MAENGELLISTE", ln=True)
    pdf.ln(5)
    
    # --- INHALT ---
    for i, item in enumerate(data):
        if pdf.get_y() > 220: pdf.add_page()
        
        pdf.set_font("Arial", 'B', 12); pdf.set_text_color(204, 0, 0)
        kat = make_safe_text(item.get('kategorie', 'Mangel'))
        prio = make_safe_text(item.get('prioritaet', 'Mittel'))
        titel = f"{i+1}. {kat} ({prio})"
        pdf.cell(0, 8, titel, ln=True)
        
        pdf.set_font("Arial", '', 10); pdf.set_text_color(0,0,0)
        pdf.multi_cell(0, 5, f"Mangel: {make_safe_text(item.get('mangel', '-'))}")
        pdf.ln(2)
        pdf.multi_cell(0, 5, f"Verstoss: {make_safe_text(item.get('verstoss', '-'))}")
        pdf.ln(2)
        pdf.multi_cell(0, 5, f"Massnahme: {make_safe_text(item.get('massnahme', '-'))}")
        pdf.ln(5)
        
        # Bild
        img_path = None
        temp_created = False
        if m_type == "video":
            frame = extract_frame(m_files[0], item.get('zeitstempel_sekunden', 0))
            if frame is not None:
                img_path = f"temp_{i}.jpg"
                cv2.imwrite(img_path, cv2.cvtColor(frame, cv2.COLOR_RGB2BGR))
                temp_created = True
        elif m_type == "images":
            idx = item.get('bild_index', 0)
            if idx < len(m_files): img_path = m_files[idx]
            
        if img_path:
            try: pdf.image(img_path, x=20, w=120)
            except: pass
            pdf.ln(10)
            if temp_created and os.path.exists(img_path): os.remove(img_path)
            
    # --- FOOTER ---
    if pdf.get_y() > 200: pdf.add_page()
    pdf.ln(15)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "4. FREIGABE", ln=True)
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 10, "Dieser Bericht wurde generiert durch SafeSite Drohne.", ln=True)
    pdf.set_font("Arial", 'I', 9)
    pdf.multi_cell(0, 5, "Hinweis: Dieser Bericht dient als visuelle Unterstuetzung. Er entbindet die zustaendige Bauleitung nicht von der gesetzlichen Kontrollpflicht.")
    pdf.ln(20)
    
    # Angepasste Breiten
    w_label = 40
    w_name = 65 
    
    pdf.set_font("Arial", 'B', 11); pdf.set_text_color(0,0,0)
    pdf.cell(w_label, 10, "Erstellt durch:", ln=0)
    pdf.set_font("Arial", '', 11)
    pdf.cell(w_name, 10, make_safe_text(f"{inspektor}"), ln=0)
    pdf.cell(0, 10, "_______________________ (Datum/Unterschrift)", ln=True, align='R')
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(w_label, 10, "Verantwortlicher:", ln=0)
    pdf.set_font("Arial", '', 11)
    pdf.cell(w_name, 10, "Bauleitung / Polier", ln=0)
    pdf.cell(0, 10, "_______________________ (Datum/Unterschrift)", ln=True, align='R')

    out = "Bericht.pdf"
    pdf.output(out)
    return out

# --- WORD GENERATOR ---
def create_word(data, m_type, m_files, projekt, inspektor, status):
    if not WORD_AVAILABLE: return None
    doc = Document()
    
    if os.path.exists(LOGO_FILE):
        try:
            doc.add_picture(LOGO_FILE, width=Inches(1.5))
            doc.paragraphs[-1].alignment = 2 
        except: pass

    doc.add_heading('SICHERHEITS-INSPEKTION (DROHNE)', 0)
    p = doc.add_paragraph()
    p.add_run("Projekt: ").bold = True; p.add_run(f"{projekt}\n")
    p.add_run("Datum: ").bold = True; p.add_run(f"{date.today().strftime('%d.%m.%Y')}\n")
    p.add_run("Inspektor: ").bold = True; p.add_run(f"{inspektor}\n")
    p.add_run("Status: ").bold = True; p.add_run(f"{status}") 

    doc.add_heading('1. ZUSAMMENFASSUNG / MÄNGEL', level=1)
    
    for i, item in enumerate(data):
        doc.add_heading(f"{i+1}. {item.get('kategorie', 'Mangel')}", level=2)
        p = doc.add_paragraph()
        p.add_run("Mangel: ").bold = True; p.add_run(f"{item.get('mangel')}\n")
        p.add_run("Verstoss: ").bold = True; p.add_run(f"{item.get('verstoss')}\n")
        p.add_run("Massnahme: ").bold = True; p.add_run(f"{item.get('massnahme')}")
        
        img_path = None
        temp_created = False
        if m_type == "video":
            frame = extract_frame(m_files[0], item.get('zeitstempel_sekunden', 0))
            if frame is not None:
                img_path = f"temp_word_{i}.jpg"; cv2.imwrite(img_path, cv2.cvtColor(frame, cv2.COLOR_RGB2BGR)); temp_created = True
        elif m_type == "images":
            idx = item.get('bild_index', 0)
            if idx < len(m_files): img_path = m_files[idx]
        
        if img_path:
            try: doc.add_picture(img_path, width=Inches(4.5))
            except: pass
            if temp_created and os.path.exists(img_path): os.remove(img_path)

    doc.add_page_break()
    doc.add_heading('4. FREIGABE', level=1)
    doc.add_paragraph("Dieser Bericht wurde generiert durch SafeSite Drohne.")
    p = doc.add_paragraph("Hinweis: Dient als visuelle Unterstützung.")
    p.italic = True
    doc.add_paragraph(f"\nErstellt durch: {inspektor} \t____________________")
    doc.add_paragraph(f"\nVerantwortlicher: \t\t\t____________________")

    out = "Bericht.docx"
    doc.save(out)
    return out

# ==========================================
# 2. APP OBERFLÄCHE
# ==========================================
if 'app_step' not in st.session_state: st.session_state.app_step = 'screen_a'
if 'analysis_data' not in st.session_state: st.session_state.analysis_data = []
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'current_page' not in st.session_state: st.session_state.current_page = 'home'
if 'username' not in st.session_state: st.session_state.username = None
if 'show_login' not in st.session_state: st.session_state.show_login = False

# SIDEBAR
with st.sidebar:
    if os.path.exists(LOGO_FILE):
        st.image(LOGO_FILE, use_container_width=True)
        
    st.title("Menü")
    page_options = ["🏠 Startseite", "🔍 SafeSite-Check", "📋 SUVA Regeln", "⚖️ BauAV", "🚨 Notfallmanagement", "🧪 Gefahrstoffkataster", "🌤️ Wetter-Warnungen"]
    p_map = {'home':0, 'safesite':1, 'suva':2, 'bauav':3, 'notfall':4, 'gefahrstoff':5, 'wetter':6, 'kunden':7}
    
    # Admin-Menüpunkt hinzufügen, wenn Admin eingeloggt
    if is_admin():
        page_options.append("👥 Kundenverwaltung")
        p_map['kunden'] = len(page_options) - 1
    
    curr_idx = p_map.get(st.session_state.current_page, 0)
    # Sicherstellen, dass der Index nicht außerhalb des Bereichs liegt
    if curr_idx >= len(page_options):
        curr_idx = 0
        st.session_state.current_page = 'home'
    
    page = st.radio("Bereich wählen:", page_options, index=curr_idx)
    
    if page == "🏠 Startseite": st.session_state.current_page = 'home'
    elif page == "🔍 SafeSite-Check": st.session_state.current_page = 'safesite'
    elif page == "📋 SUVA Regeln": st.session_state.current_page = 'suva'
    elif page == "⚖️ BauAV": st.session_state.current_page = 'bauav'
    elif page == "🚨 Notfallmanagement": st.session_state.current_page = 'notfall'
    elif page == "🧪 Gefahrstoffkataster": st.session_state.current_page = 'gefahrstoff'
    elif page == "🌤️ Wetter-Warnungen": st.session_state.current_page = 'wetter'
    elif page == "👥 Kundenverwaltung": st.session_state.current_page = 'kunden'
    
    st.divider()
    
    # Login-Bereich in Sidebar
    if not st.session_state.logged_in:
        st.markdown("### 🔐 Login")
        
        with st.form("login_form", clear_on_submit=False):
            u = st.text_input("Username / Email", key="sidebar_username")
            p = st.text_input("Passwort", type="password", key="sidebar_password")
            if st.form_submit_button("Einloggen", use_container_width=True, type="primary"):
                users = load_users()
                # Direkter Login-Check (für Admin oder wenn genau der Key existiert)
                if u in users and users[u] == p:
                    st.session_state.logged_in = True
                    st.session_state.username = u
                    st.session_state.show_login = False
                    st.rerun()
                else:
                    # Prüfe ob es ein Kunde ist (Email oder Benutzername)
                    kunde_id, kunde_data = get_customer_by_username_or_email(u)
                    if kunde_data:
                        # Finde den korrekten Login-Key (kann Email oder Username sein)
                        customer_email = kunde_data.get('email', '')
                        customer_username = kunde_data.get('username', '')
                        
                        # Prüfe ob Login mit Email existiert
                        if customer_email and customer_email in users and users[customer_email] == p:
                            st.session_state.logged_in = True
                            st.session_state.username = customer_email  # Speichere Email als username
                            st.session_state.show_login = False
                            st.rerun()
                        # Prüfe ob Login mit Benutzername existiert
                        elif customer_username and customer_username in users and users[customer_username] == p:
                            st.session_state.logged_in = True
                            st.session_state.username = customer_username
                            st.session_state.show_login = False
                            st.rerun()
                    
                    st.error("❌ Falscher Username/Email oder Passwort!")
    else:
        st.markdown("### 👤 Benutzer")
        st.info(f"✅ Eingeloggt als: **{st.session_state.username}**")
        
        # Credits-Anzeige nur für Kunden (nicht Admin)
        if not is_admin() and st.session_state.username:
            credits = get_customer_credits(st.session_state.username)
            st.metric("🪙 SafeSite Credits", credits)
        
        if st.button("Logout", use_container_width=True): 
            st.session_state.logged_in = False
            st.session_state.username = None
            st.session_state.current_page = 'home'
            st.session_state.app_step = 'screen_a'
            st.session_state.analysis_data = []
            st.rerun()

# HAUPTBEREICH
if st.session_state.current_page == 'home':
    if os.path.exists(TITELBILD_FILE):
        st.image(TITELBILD_FILE, use_container_width=True)
    
    st.header("🏠 Willkommen bei SafeSite Drohne")
    
    st.markdown("## Sicherheit, die sich auszahlt.")
    st.markdown("""
    SafeSite Drohne ist mehr als nur eine Kamera in der Luft. Wir liefern Ihnen ein komplettes System zur Unfallprävention und Dokumentation – entwickelt von Polieren für den täglichen Einsatz.
    """)
    
    st.markdown("---")
    
    # Menüpunkte als Buttons
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔍 SafeSite-Check", use_container_width=True, type="primary"):
            st.session_state.current_page = 'safesite'
            st.rerun()
        
        if st.button("📋 SUVA Regeln", use_container_width=True):
            st.session_state.current_page = 'suva'
            st.rerun()
        
        if st.button("🚨 Notfallmanagement", use_container_width=True):
            st.session_state.current_page = 'notfall'
            st.rerun()
        
        if st.button("🧪 Gefahrstoffkataster", use_container_width=True):
            st.session_state.current_page = 'gefahrstoff'
            st.rerun()
        
        if st.button("🌤️ Wetter-Warnungen", use_container_width=True):
            st.session_state.current_page = 'wetter'
            st.rerun()
    
    with col2:
        if st.button("⚖️ BauAV", use_container_width=True):
            st.session_state.current_page = 'bauav'
            st.rerun()
        
        # Kundenverwaltung nur für Admin
        if is_admin():
            if st.button("👥 Kundenverwaltung", use_container_width=True):
                st.session_state.current_page = 'kunden'
                st.rerun()

elif st.session_state.current_page == 'safesite':
    if not st.session_state.logged_in:
        st.header("🔍 SafeSite-Check")
        st.warning("⚠️ Sie müssen sich anmelden, um den SafeSite-Check zu verwenden.")
    else:
        # APP START
        if st.session_state.app_step == 'screen_a':
            st.subheader("Neuer Auftrag")
            mode = st.radio("Quelle:", ["📹 Video", "📸 Fotos"], horizontal=True)
            files = []
            
            if mode == "📹 Video":
                st.info("💡 **Hinweis:** Auf mobilen Geräten wählen Sie bitte Videos über den Datei-Explorer aus.")
                vf = st.file_uploader("Video hochladen", type=["mp4", "mov", "avi"], help="Unterstützte Formate: MP4, MOV, AVI")
                if vf:
                    try:
                        st.success(f"✅ Video ausgewählt: {vf.name}")
                        if st.button("Analyse starten", type="primary", use_container_width=True):
                            try:
                                suffix = os.path.splitext(vf.name)[1] if os.path.splitext(vf.name)[1] else '.mp4'
                                t = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                                # Datei in Chunks lesen für bessere Performance
                                chunk_size = 8192
                                while True:
                                    chunk = vf.read(chunk_size)
                                    if not chunk:
                                        break
                                    t.write(chunk)
                                t.close()
                                files.append(t.name)
                                st.session_state.m_type = "video"
                                st.session_state.m_files = files
                                st.session_state.app_step = 'screen_b'
                                st.rerun()
                            except Exception as e:
                                st.error(f"❌ Fehler beim Hochladen: {str(e)}")
                                st.info("💡 Bitte versuchen Sie es erneut oder verwenden Sie eine kleinere Datei.")
                    except Exception as e:
                        st.error(f"❌ Fehler beim Lesen der Datei: {str(e)}")
            else:
                st.info("💡 **Hinweis:** Auf mobilen Geräten wählen Sie bitte Fotos über den Datei-Explorer aus.")
                pf = st.file_uploader(
                    "Fotos hochladen", 
                    type=["jpg", "jpeg", "png", "heic", "heif", "webp"], 
                    accept_multiple_files=True,
                    help="Wählen Sie ein oder mehrere Fotos aus. Unterstützt: JPG, PNG, HEIC (iPhone), WEBP"
                )
                if pf:
                    try:
                        st.success(f"✅ {len(pf)} Foto(s) ausgewählt")
                        # Zeige Dateinamen an
                        for idx, f in enumerate(pf[:5]):  # Zeige max. 5 Dateien
                            st.caption(f"📷 {f.name}")
                        if len(pf) > 5:
                            st.caption(f"... und {len(pf) - 5} weitere")
                        
                        if st.button("Analyse starten", type="primary", use_container_width=True):
                            with st.spinner("Bilder werden verarbeitet..."):
                                try:
                                    for f in pf:
                                        # Original-Dateiendung beibehalten
                                        original_ext = os.path.splitext(f.name)[1].lower() if os.path.splitext(f.name)[1] else ''
                                        # Verwende passende Endung basierend auf Dateityp
                                        if original_ext in ['.heic', '.heif']:
                                            suffix = '.heic'  # Wird später konvertiert
                                        elif original_ext in ['.jpg', '.jpeg']:
                                            suffix = '.jpg'
                                        elif original_ext == '.png':
                                            suffix = '.png'
                                        elif original_ext == '.webp':
                                            suffix = '.webp'
                                        else:
                                            suffix = '.jpg'  # Standard
                                        
                                        t = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                                        # Datei in Chunks lesen für bessere Performance und Fehlerbehandlung
                                        chunk_size = 8192
                                        while True:
                                            chunk = f.read(chunk_size)
                                            if not chunk:
                                                break
                                            t.write(chunk)
                                        t.close()
                                        files.append(t.name)
                                    
                                    # Konvertiere HEIC/HEIF Dateien falls nötig
                                    converted_files = []
                                    for f_path in files:
                                        converted_path = convert_image_if_needed(f_path)
                                        converted_files.append(converted_path)
                                    
                                    st.session_state.m_type = "images"
                                    st.session_state.m_files = converted_files
                                    st.session_state.app_step = 'screen_b'
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"❌ Fehler beim Hochladen der Dateien: {str(e)}")
                                    st.info("💡 Bitte versuchen Sie es erneut. Stellen Sie sicher, dass Sie eine stabile Internetverbindung haben.")
                                    # Aufräumen: Temporäre Dateien löschen
                                    for f_path in files:
                                        try:
                                            if os.path.exists(f_path):
                                                os.remove(f_path)
                                        except:
                                            pass
                    except Exception as e:
                        st.error(f"❌ Fehler beim Lesen der Dateien: {str(e)}")
                        st.info("💡 Bitte versuchen Sie es erneut oder wählen Sie andere Dateien aus.")

        elif st.session_state.app_step == 'screen_b':
            st.subheader("🕵️‍♂️ KI-Analyse (Gemini 3.0)")
            if st.session_state.m_type == "video": st.video(st.session_state.m_files[0])
            else: 
                cols = st.columns(3)
                for i, f in enumerate(st.session_state.m_files):
                    with cols[i % 3]: st.image(f, caption=f"Bild {i+1}")

            if not st.session_state.analysis_data:
                # Verbesserter Ladebalken mit progressiven Nachrichten
                progress_placeholder = st.empty()
                status_placeholder = st.empty()
                
                progress_messages = [
                    "🔍 SafeSite analysiert Gefahren...",
                    "🔍 SafeSite prüft Absturzsicherungen...",
                    "🔍 SafeSite kontrolliert Gerüste...",
                    "🔍 SafeSite überprüft Schweizer Normen...",
                    "🔍 SafeSite bewertet Sicherheitsrisiken...",
                    "🔍 SafeSite erstellt Analyse..."
                ]
                
                try:
                    genai.configure(api_key=API_KEY)
                    
                    # Detaillierter Prompt mit spezifischen Schweizer Normen
                    prompt = f"""
Du bist ein äusserst strenger und erfahrener Schweizer Bau-Sicherheitsprüfer (SiBe) mit tiefem Wissen der BauAV und SUVA-Richtlinien.

KRITISCH: Du erhältst {len(st.session_state.m_files)} Bilder zur Analyse. Analysiere JEDES Bild MILLIMETERGENAU und SYSTEMATISCH nach ALLEN relevanten Schweizer Sicherheitsnormen (BauAV und SUVA). Es ist deine PFLICHT, JEDEN noch so kleinen Verstoss zu erkennen!

KRITISCH: Der Parameter "bild_index" muss der Index des Bildes sein (0 für das erste Bild, 1 für das zweite Bild, etc.). Jeder Mangel muss dem korrekten bild_index zugeordnet werden!

PRÜFUNGSPROTOKOLL - Du musst ALLE folgenden Punkte für JEDES Bild systematisch durchgehen:

1. GERÜSTE (BauAV Art. 47, 48, 49):
   - Abstand Gerüst-Fassade: GENAU < 30cm? (Jeder cm darüber ist ein Mangel!)
   - Dreiteiliger Seitenschutz: Holm (oberste Querleiste), Zwischenholm (mittlere Querleiste), Bordbrett (untere Abschlussleiste) ALLE vorhanden?
   - Seitenschutz-Höhe: Mindestens 1.0m? Gemessen vom Laufsteg aus?
   - Beläge: Lückenlos verlegt? Keine Löcher? Keine Überhänge > 10cm?
   - Belag-Tragfähigkeit: Mind. 2.0 kN/m²? (Visuell: solide Beläge?)
   - Gerüstständer: Auf standsicherem Untergrund? Ausreichend Abstützungen?
   - Verbindungen: Alle Steckverbindungen korrekt? Keine lockeren Teile?
   - Zugänge: Treppen oder Leitern vorhanden? Sicher befestigt?
   - Absturzsicherung: An allen offenen Seiten vorhanden?
   - Tägliche Kontrolle: Kontrollschild sichtbar? (Wenn nicht, als Mangel melden)

2. ABSTURZKANTEN (BauAV Art. 17, 18):
   - Seitenschutz ab 2.0m Absturzhöhe: MUSS vorhanden sein!
   - Seitenschutz-Höhe: Mindestens 1.0m hoch?
   - Seitenschutz-Festigkeit: Stabil genug? (Mind. 1.0 kN/m horizontal)
   - Dachkanten ab 3.0m: Zusätzliche Sicherung vorhanden? (Netze, Seile, etc.)
   - Absturzkanten deutlich sichtbar? (Markierungen vorhanden?)
   - Keine provisorischen Sicherungen? (Nur genormte Systeme!)

3. BODENÖFFNUNGEN (BauAV Art. 19):
   - Durchbruchsichere Abdeckung: Vorhanden? (Mind. 5.0 kN/m² Tragfähigkeit)
   - Abdeckung gesichert: Gegen Verschieben gesichert? (Schrauben, Klammern, etc.)
   - Öffnung markiert: Warnzeichen oder Absperrung vorhanden?
   - Abdeckung vollständig: Keine Löcher oder Risse?
   - Grösse der Öffnung: > 20cm? Dann Sicherung PFLICHT!

4. GRÄBEN UND SCHÄCHTE (BauAV Art. 20, 21):
   - Verspriesst ab 1.50m Tiefe: MUSS vorhanden sein! (Bei <1.50m nur wenn Böschung nicht möglich)
   - Böschungswinkel: Max. 45° bei bindigen Böden, max. 63° bei nicht-bindigen Böden?
   - Verspriesst: Ausreichend stabil? (Verbindungsmittel vorhanden?)
   - Durchbruchsicherheit: Verspriesst stützt Wände ausreichend?
   - Bei fliessenden Böden: Sicherung bereits ab 1.0m Tiefe?
   - Absturzsicherung: Geländer oder Abdeckung am Grabenrand?
   - Zugänge: Sichere Treppen oder Leitern vorhanden?
   - Wasserhaltung: Bei Wasseransammlung korrekt abgepumpt/gesichert?

5. LEITERN (BauAV Art. 23, 24):
   - Verwendungszweck: Nur für kurzzeitige Arbeiten (< 2h)? (Wenn länger, Mangel!)
   - Anlehnwinkel: 65-75°? (Zu steil oder zu flach = Mangel!)
   - Überragung: Mindestens 1.0m über Austrittsstelle?
   - Sicherung: Gegen Wegrutschen gesichert? (Oben und unten)
   - Zustand: Keine defekten Sprossen? Keine Risse?
   - Standsicherheit: Auf festem, ebenem Untergrund?
   - Anschlagsicherheit: Oberer Teil fest verbunden?

6. PERSÖNLICHE SCHUTZAUSRÜSTUNG (PSA) - BauAV Art. 6, 7, SUVA:
   - Schutzhelm (BauAV Art. 6): MUSS getragen werden bei:
     * Hochbau- und Brückenbauarbeiten bis Rohbauabschluss
     * Arbeiten im Bereich von Kranen, Aushubgeräten, Spezialtiefbaumaschinen
     * Graben- und Schachtbau, Baugruben
     * Gerüstbauarbeiten
     * Rückbau/Abbrucharbeiten
     * Arbeiten an/in Rohrleitungen
     * Bei Gefahr durch herunterfallende Gegenstände
   - Schutzhelm MIT KINNBAND (BauAV Art. 6 Abs. 3): Bei Seilsicherung, Arbeiten am hängenden Seil, Helikopter-Bereich
   - Warnkleider (BauAV Art. 7): Bei Verkehrsmitteln (Baumaschinen, Transportfahrzeuge) oder öffentlichen Verkehrswegen
   - Sicherheitsschuhe: Getragen? (PFLICHT!)
   - Schutzbrille: Bei Staub, Spritzern, Splittern getragen?
   - Gehörschutz: Bei Lärm > 85 dB(A) getragen?
   - Schutzhandschuhe: Bei scharfen Kanten, Chemikalien getragen?
   - Atemschutz: Bei Staub, Dämpfen, Gasen getragen?
   - PSA korrekt angelegt? (Helm richtig aufgesetzt, nicht nur auf dem Kopf?)

7. VERKEHRSWEGE (BauAV Art. 25, SUVA):
   - Breite: Mindestens 0.80m frei? (Für Fussgänger)
   - Hindernisse: Weg vollständig frei?
   - Stolperstellen: Keine Kabel, Werkzeuge, Materialien auf dem Weg?
   - Beleuchtung: Bei schlechter Sicht (Dämmerung, Innenräume) vorhanden?
   - Markierung: Wege deutlich gekennzeichnet?
   - Gefälle: Maximal 10%? (Bei >10% Treppen oder Rampen)
   - Rutschfest: Belag rutschfest? (Keine glatten, nassen Flächen ohne Profil?)

8. MATERIALLAGERUNG (BauAV Art. 26, SUVA):
   - Stapelung: Material stabil gestapelt? (Nicht höher als 3x Basisbreite)
   - Gänge: Mindestens 0.8m freigehalten zwischen Stapeln?
   - Umkippgefahr: Keine instabilen Stapel? (Kippen erkennbar?)
   - Schweres Material: Unten gelagert? (Schweres zuunterst)
   - Gefahrstoffe: Korrekt gekennzeichnet? (Wenn sichtbar)
   - Lagerhöhe: Nicht zu hoch? (Max. Arbeitshöhe + 1m)
   - Standsicherheit: Stapel nicht zu nah an Absturzkanten?

9. ARBEITEN IN HÖHE (BauAV Art. 27, SUVA):
   - Bereich unterhalb: Abgesperrt oder Schutzdächer vorhanden?
   - Werkzeuge: Gegen Herunterfallen gesichert? (Leinen, Taschen, etc.)
   - Material: Alles gegen Herunterfallen gesichert?
   - Personen im Gefahrenbereich: Niemand unter Arbeiten in Höhe?
   - Gerüste/Arbeitsbühnen: Korrekt aufgestellt und gesichert?

10. KRANE UND LASTEN (BauAV Art. 28, 29, SUVA):
    - Personen unter Lasten: ABSOLUT VERBOTEN! (Kritischer Mangel!)
    - Anschlagmittel: Korrekt? (4-fache Sicherheit bei Stahlseilen)
    - Lastgewicht: Nicht über Nennlast des Krans?
    - Kommunikation: Einweiser vorhanden? Sichtkontakt zum Kranführer?
    - Kranstellung: Auf festem Untergrund? Ausreichend Abstützungen?
    - Seilführungen: Korrekt? Keine Überlastungen erkennbar?
    - Warnsignale: Hörbar? (Hupen, etc.)

11. ELEKTRIZITÄT (BauAV Art. 30, SUVA):
    - Kabel: Nicht beschädigt? (Keine offenen Stellen, Isolierung intakt?)
    - Kabel: Nicht im Weg? (Nicht auf Boden liegend wo gefahren wird?)
    - Steckdosen: Abgedeckt? (Wasserdicht bei Aussenbereich?)
    - FI-Schutzschalter: Vorhanden? (Bei sichtbaren Anschlüssen)
    - Hochspannung: Ausreichend Abstand? (Bei Freileitungen)

12. BRANDSCHUTZ (BauAV Art. 31, SUVA):
    - Fluchtwege: Frei? (Keine Blockierungen?)
    - Löschmittel: Sichtbar? (Feuerlöscher, etc. wenn erkennbar)
    - Brandlasten: Minimiert? (Keine unnötigen brennbaren Materialien)

13. LÄRM UND VIBRATIONEN (SUVA):
    - Gehörschutz: Getragen bei lauten Maschinen?
    - Warnschilder: Bei Lärmbereichen vorhanden?

14. STAUB UND GEFÄHRSTOFFE (BauAV Art. 32, SUVA):
    - Atemschutz: Getragen bei sichtbarem Staub?
    - Absaugung: Vorhanden bei staubigen Arbeiten?
    - Gefahrstoffkennzeichnung: Sichtbar? (Wenn Gefahrstoffe erkennbar)

15. BAUFAHRZEUGE (SUVA):
    - Rückfahrwarnsystem: Funktioniert? (Piepser hörbar?)
    - Toter Winkel: Einweiser vorhanden?
    - Geschwindigkeit: Angemessen? (Max. 10 km/h auf Baustelle)

ABSOLUT KRITISCHE REGELN (BauAV SR-832.311.141):
- DU MUSST ALLES prüfen gemäss Bauarbeitenverordnung (BauAV) und SUVA-Richtlinien! Auch wenn etwas "vielleicht ok aussieht", prüfe es MILLIMETERGENAU!
- KEINE "Ist ok" Bewertungen ohne detaillierte Prüfung ALLER Kriterien aus der BauAV!
- Jeder noch so kleine Verstoss gegen BauAV oder SUVA-Regeln MUSS als Mangel erkannt werden!
- Wenn du etwas NICHT SICHER ERKENNEN kannst, ist das ein Mangel! ("Unklar, ob Seitenschutz korrekt montiert gemäss BauAV Art. 22" = Mangel!)
- Referenziere IMMER die genauen BauAV-Artikel (z.B. "BauAV Art. 6 Abs. 2", "BauAV Art. 17", "BauAV Art. 20")
- Priorität: 
  * "Kritisch" = Lebensgefahr (z.B. kein Helm gemäss BauAV Art. 6, Person unter Last gemäss BauAV Art. 28, Absturzgefahr >2m ohne Schutz gemäss BauAV Art. 17)
  * "Hoch" = Schwere Verstösse (z.B. Gerüst ohne Seitenschutz gemäss BauAV Art. 22, Graben >1.5m ohne Verspriesst gemäss BauAV Art. 20)
  * "Mittel" = Normative Abweichungen (z.B. Abstand Gerüst-Fassade 35cm statt <30cm gemäss BauAV Art. 47)
- Analysiere JEDES Bild separat und setze den bild_index korrekt (0, 1, 2, etc. je nach Bildnummer)
- Wenn du mehrere Mängel in einem Bild siehst, erstelle für JEDEN einen separaten Eintrag mit korrekter BauAV-Referenz!
- Beachte: Diese Verordnung (BauAV SR-832.311.141) ist bindend - alle Vorschriften MÜSSEN eingehalten werden!

WICHTIG FÜR DIE BERICHTGESTALTUNG:
- Schreibe DETAILLIERTE, PROFESSIONELLE Texte wie ein erfahrener SiBe!
- Jeder Mangel muss eine präzise Beschreibung haben mit konkreten Angaben (z.B. "Abstand Gerüst-Fassade 50cm statt <30cm", "Böschungswinkel ca. 60° statt max. 45°")
- Beschreibe den BEFUND genau (was siehst du?), die NORM (welche Vorschrift wird verletzt?), und die MASSNAHME (was muss gemacht werden?)
- Verwende präzise Fachbegriffe und konkrete Messwerte wo möglich
- Strukturiere die Mängel logisch (z.B. "Absturzsicherung an der Baugrubenkante", "Fassadengerüst – Innenliegender Absturz")
- Bei kritischen Mängeln: Beschreibe die AKUTE GEFAHR und warum SOFORTMASSNAHMEN erforderlich sind
- Berücksichtige Witterungsverhältnisse (Schnee, Eis, Regen) wenn sichtbar
- Beschreibe auch Situationen, die du NICHT SICHER ERKENNEN kannst (z.B. "Aufgrund der Distanz schwer erkennbar, aber...")

Antworte NUR als JSON Liste:
[{{"kategorie": "...", "prioritaet": "Kritisch/Hoch/Mittel", "mangel": "DETAILLIERTE BESCHREIBUNG des Mangels mit konkreten Angaben (z.B. 'An fast allen Baugrubenrändern (besonders im Bereich des noch nicht hinterfüllten Kellers im rechten Bildteil) fehlt der vorgeschriebene Seitenschutz. Es besteht unmittelbare Lebensgefahr durch Absturz in die Grube.')", "verstoss": "GENAUER Verstoss mit Artikel-Referenz (z.B. 'Verstoss BauAV Art. 17 - Bei Absturzhöhen über 2m ist ein dreiteiliger Seitenschutz zwingend')", "massnahme": "KONKRETE, AUSFÜHRLICHE Massnahme (z.B. 'Sofortige Absperrung (mind. 1.5m - 2m Abstand zur Kante) oder Montage eines festen Geländers. Geologen/Geotechniker hinzuziehen. Böschungswinkel kontrollieren.')", "zeitstempel_sekunden": 0, "bild_index": 0}}]

Beispiel für eine professionelle Mangelbeschreibung:
{{"kategorie": "Baugruben und Erdarbeiten", "prioritaet": "Kritisch", "mangel": "An fast allen Baugrubenrändern (besonders im Bereich des noch nicht hinterfüllten Kellers im rechten Bildteil) fehlt der vorgeschriebene Seitenschutz. Es besteht unmittelbare Lebensgefahr durch Absturz in die Grube. Die Böschungen sind steil und mit Schnee bedeckt. Durch Schmelzwasser besteht akute Rutschgefahr.", "verstoss": "Verstoss BauAV Art. 17 - Bei Absturzhöhen über 2m ist ein dreiteiliger Seitenschutz zwingend. Verstoss BauAV Art. 59 - Böschungswinkel zu steil.", "massnahme": "Sofortige Absperrung (mind. 1.5m - 2m Abstand zur Kante) oder Montage eines festen Geländers. Geologen/Geotechniker hinzuziehen. Böschungswinkel kontrollieren. Bei aufgeweichtem Boden Böschung abflachen oder verbauen.", "zeitstempel_sekunden": 0, "bild_index": 0}}
"""
                    
                    # --- HIER IST DIE SCHLAUE SCHLEIFE ---
                    # Wir probieren die Modelle der Reihe nach durch.
                    # Wenn 3.0 nicht geht, nimmt er automatisch 2.0 oder 1.5
                    model_names = [
                        'gemini-3-pro-preview', 
                        'gemini-2.0-flash-exp', 
                        'gemini-1.5-pro',
                        'gemini-1.5-flash'
                    ]
                    
                    found_result = False
                    
                    # Progress-Tracker
                    progress_step = 0
                    start_time = time.time()
                    
                    for mn in model_names:
                        try:
                            # Update Progress Message
                            if progress_step < len(progress_messages):
                                status_placeholder.info(f"🔄 {progress_messages[progress_step]}")
                                progress_step += 1
                            
                            model = genai.GenerativeModel(mn)
                            if st.session_state.m_type == "video":
                                status_placeholder.info("🔄 SafeSite lädt Video hoch...")
                                f = genai.upload_file(st.session_state.m_files[0])
                                # Warten (Fix für Hänger) mit Progress
                                while f.state.name == "PROCESSING":
                                    elapsed = int(time.time() - start_time)
                                    status_placeholder.info(f"🔄 SafeSite verarbeitet Video... ({elapsed}s)")
                                    time.sleep(2)
                                    f = genai.get_file(f.name)
                                
                                status_placeholder.info("🔄 SafeSite analysiert Video nach Schweizer Normen...")
                                res = model.generate_content([f, prompt], generation_config={"response_mime_type": "application/json"})
                            else:
                                status_placeholder.info("🔄 SafeSite lädt Bilder...")
                                # Öffne Bilder und konvertiere bei Bedarf
                                imgs = []
                                for idx, p in enumerate(st.session_state.m_files):
                                    try:
                                        status_placeholder.info(f"🔄 SafeSite verarbeitet Bild {idx+1}/{len(st.session_state.m_files)}...")
                                        img = Image.open(p)
                                        # Stelle sicher, dass Bild im RGB-Format ist
                                        if img.mode != 'RGB':
                                            img = img.convert('RGB')
                                        imgs.append(img)
                                    except Exception as e:
                                        st.warning(f"⚠️ Fehler beim Öffnen von {os.path.basename(p)}: {str(e)}")
                                        # Versuche mit cv2 als Fallback
                                        try:
                                            img_array = cv2.imread(p)
                                            if img_array is not None:
                                                img_rgb = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
                                                img = Image.fromarray(img_rgb)
                                                imgs.append(img)
                                        except:
                                            st.error(f"❌ Konnte Bild {os.path.basename(p)} nicht verarbeiten")
                                
                                if not imgs:
                                    status_placeholder.error("❌ Keine Bilder konnten verarbeitet werden.")
                                    st.error("❌ Keine Bilder konnten verarbeitet werden. Bitte versuchen Sie andere Dateiformate.")
                                    continue
                                
                                status_placeholder.info("🔄 SafeSite analysiert Bilder nach Schweizer Normen (BauAV & SUVA)...")
                                # Zeige Progress während der Analyse
                                elapsed = int(time.time() - start_time)
                                status_placeholder.info(f"🔄 SafeSite prüft Gerüste, Absturzkanten, Gräben... ({elapsed}s)")
                                
                                res = model.generate_content([prompt] + imgs, generation_config={"response_mime_type": "application/json"})
                            
                            # Analyse abgeschlossen
                            elapsed = int(time.time() - start_time)
                            status_placeholder.success(f"✅ SafeSite Analyse abgeschlossen! ({elapsed}s)")
                            time.sleep(0.5)  # Kurze Pause, damit die Erfolgsmeldung sichtbar ist
                            
                            # Wenn wir hier sind, hat es geklappt!
                            st.session_state.analysis_data = json.loads(clean_json(res.text))
                            found_result = True
                            break # Schleife beenden, wir haben ein Ergebnis
                        except Exception as e:
                            elapsed = int(time.time() - start_time)
                            status_placeholder.warning(f"⚠️ Versuche nächstes Modell... ({elapsed}s)")
                            continue # Fehler beim Modell? Nächstes probieren!
                    
                    # Aufräumen der Placeholders
                    progress_placeholder.empty()
                    
                    if not found_result:
                        status_placeholder.error("❌ Alle KI-Modelle sind gerade ausgelastet oder nicht erreichbar. Bitte später versuchen.")
                    else:
                        status_placeholder.empty()  # Entferne Status-Nachricht
                        st.rerun()
                        
                except Exception as e:
                    progress_placeholder.empty()
                    status_placeholder.error(f"❌ Fehler: {e}")
                    st.error(f"Fehler: {e}")

            if st.session_state.analysis_data:
                st.success(f"⚠️ {len(st.session_state.analysis_data)} Mängel gefunden")
                
                # Credits-Anzeige für Kunden (nicht Admin)
                if not is_admin() and st.session_state.username:
                    credits = get_customer_credits(st.session_state.username)
                    col_credits = st.columns([2, 1])
                    with col_credits[1]:
                        if credits < 1:
                            st.error(f"🪙 Credits: {credits} (Nicht genügend für Bericht!)")
                        else:
                            st.info(f"🪙 Verbleibende Credits: **{credits}**")
                    st.divider()
                
                st.markdown("### 📝 Projektdaten für Bericht")
                c_a, c_b = st.columns(2)
                with c_a:
                    proj = st.text_input("Projektname", value="Überbauung 'Luegisland', Wohlen AG")
                    insp = st.text_input("Inspektor Name", value="Dominik Marti")
                with c_b:
                    stat = st.selectbox("Status", ["⚠️ Massnahmen erforderlich", "✅ In Ordnung", "🛑 Kritisch - Baustopp"])
                st.divider()

                with st.form("check"):
                    confirmed = []
                    for i, item in enumerate(st.session_state.analysis_data):
                        c1, c2 = st.columns([1,3])
                        with c1:
                            if st.session_state.m_type == "video":
                                frm = extract_frame(st.session_state.m_files[0], item.get('zeitstempel_sekunden', 0))
                                if frm is not None: st.image(frm)
                            else:
                                idx = item.get('bild_index', 0)
                                if idx < len(st.session_state.m_files): st.image(st.session_state.m_files[idx])
                        with c2:
                            st.markdown(f":orange[**{item.get('prioritaet')}: {item.get('mangel')}**]")
                            st.write(item.get('massnahme'))
                            if st.checkbox("Aufnehmen", True, key=str(i)): confirmed.append(item)
                        st.divider()
                    
                    if st.form_submit_button("Berichte erstellen"):
                        # Credit-Prüfung (nur für Kunden, nicht für Admin)
                        if not is_admin():
                            username = st.session_state.username
                            credits = get_customer_credits(username)
                            if credits < 1:
                                st.error(f"⚠️ Nicht genügend Credits! Sie haben {credits} Credit(s). Bitte kontaktieren Sie den Administrator.")
                            else:
                                # Credit abbuchen
                                if deduct_credit(username):
                                    st.success(f"✅ 1 Credit abgebucht. Verbleibend: {credits - 1}")
                                    st.session_state.confirmed = confirmed
                                    st.session_state.meta_p = proj
                                    st.session_state.meta_i = insp
                                    st.session_state.meta_s = stat
                                    st.session_state.app_step = 'screen_c'
                                    st.rerun()
                                else:
                                    st.error("⚠️ Fehler beim Abziehen der Credits. Bitte versuchen Sie es erneut.")
                        else:
                            # Admin kann ohne Credits erstellen
                            st.session_state.confirmed = confirmed
                            st.session_state.meta_p = proj
                            st.session_state.meta_i = insp
                            st.session_state.meta_s = stat
                            st.session_state.app_step = 'screen_c'
                            st.rerun()

        elif st.session_state.app_step == 'screen_c':
            st.subheader("Berichte fertig!")
            
            # Credits-Anzeige nach erfolgreicher Erstellung (für Kunden)
            if not is_admin() and st.session_state.username:
                remaining_credits = get_customer_credits(st.session_state.username)
                st.info(f"🪙 Verbleibende Credits: **{remaining_credits}**")
                st.divider()
            
            p = st.session_state.get('meta_p', '')
            i = st.session_state.get('meta_i', '')
            s = st.session_state.get('meta_s', '')

            # Berichte nur einmal erstellen und im Session State speichern
            if 'pdf_file_path' not in st.session_state:
                try:
                    pdf_file = create_pdf(st.session_state.confirmed, st.session_state.m_type, st.session_state.m_files, p, i, s)
                    st.session_state.pdf_file_path = pdf_file
                except Exception as e:
                    st.error(f"❌ Fehler beim Erstellen des PDF-Berichts: {str(e)}")
                    st.stop()
            else:
                pdf_file = st.session_state.pdf_file_path
            
            if 'word_file_path' not in st.session_state and WORD_AVAILABLE:
                try:
                    word_file = create_word(st.session_state.confirmed, st.session_state.m_type, st.session_state.m_files, p, i, s)
                    if word_file:
                        st.session_state.word_file_path = word_file
                except Exception as e:
                    st.warning(f"⚠️ Fehler beim Erstellen des Word-Berichts: {str(e)}")
            
            c1, c2 = st.columns(2)
            with c1:
                try:
                    if os.path.exists(pdf_file):
                        with open(pdf_file, "rb") as f:
                            st.download_button("📄 PDF Bericht", f, "SSD_Bericht.pdf", mime="application/pdf", use_container_width=True)
                    else:
                        st.error("❌ PDF-Datei nicht gefunden")
                except Exception as e:
                    st.error(f"❌ Fehler beim Laden des PDF: {str(e)}")
            with c2:
                if WORD_AVAILABLE and 'word_file_path' in st.session_state:
                    try:
                        word_file = st.session_state.word_file_path
                        if word_file and os.path.exists(word_file):
                            with open(word_file, "rb") as f:
                                st.download_button("📝 Word Bericht", f, "SSD_Bericht.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    except Exception as e:
                        st.warning(f"⚠️ Fehler beim Laden des Word-Dokuments: {str(e)}")

            st.divider()
            st.markdown("### 📧 Versenden")
            email_to = st.text_input("Empfänger Email", placeholder="kunde@bau.ch")
            
            if email_to:
                subject = f"Sicherheitsbericht: {p}"
                body = f"Grüezi,\n\nanbei erhalten Sie den Sicherheitsbericht für das Projekt {p}.\n\nInspektor: {i}\nStatus: {s}\n\nFreundliche Grüsse\nSafeSite Drohne"
                safe_s = urllib.parse.quote(subject)
                safe_b = urllib.parse.quote(body)
                mailto = f"mailto:{email_to}?subject={safe_s}&body={safe_b}"
                
                st.link_button("📧 Email-Programm öffnen", mailto)

            if st.button("Neuer Auftrag"):
                st.session_state.app_step = 'screen_a'
                st.session_state.analysis_data = []
                # Session State für Berichte zurücksetzen
                if 'pdf_file_path' in st.session_state:
                    del st.session_state.pdf_file_path
                if 'word_file_path' in st.session_state:
                    del st.session_state.word_file_path
                st.rerun()

elif st.session_state.current_page == 'suva':
    st.header("📋 Die 8 lebenswichtigen Regeln (SUVA)")
    
    # Quellenangabe und Hinweis
    st.info("ℹ️ **Quelle:** Diese Regeln basieren auf den lebenswichtigen Regeln der SUVA. Bilder, Grafiken und Texte: © SUVA. Diese App ist keine offizielle SUVA-App, sondern dient der Arbeitssicherheit im internen Gebrauch.")
    
    st.markdown("---")
    
    suva_regeln = [
        {"titel": "1. Absturzkanten sichern", "desc": "Ab 2.0m Absturzhöhe sind Seitenschutz oder Auffangeinrichtungen zwingend.", "img": "regel_1.png"},
        {"titel": "2. Bodenöffnungen", "desc": "Jede Öffnung muss durchbruchsicher abgedeckt und fixiert sein.", "img": "regel_2.png"},
        {"titel": "3. Lasten anschlagen", "desc": "Lasten nur von instruiertem Personal anschlagen. Niemals unter schwebenden Lasten.", "img": "regel_3.png"},
        {"titel": "4. Fassadengerüste", "desc": "Ab 3.0m Absturzhöhe ist ein Fassadengerüst erforderlich.", "img": "regel_4.png"},
        {"titel": "5. Gerüstkontrolle", "desc": "Tägliche Sichtkontrolle durch den Benutzer. Beläge müssen dicht sein.", "img": "regel_5.png"},
        {"titel": "6. Sichere Zugänge", "desc": "Treppentürme sind Leitern vorzuziehen. Leitern gegen Wegrutschen sichern.", "img": "regel_6.png"},
        {"titel": "7. PSA tragen", "desc": "Helm und Sicherheitsschuhe sind Pflicht. Je nach Situation: Weste, Brille, Gehörschutz.", "img": "regel_7.png"},
        {"titel": "8. Gräben sichern", "desc": "Ab 1.50m Tiefe müssen Gräben gespriesst oder geböscht werden.", "img": "regel_8.png"}
    ]

    for r in suva_regeln:
        with st.container(border=True):
            c1, c2 = st.columns([1, 2])
            with c1:
                if os.path.exists(r["img"]):
                    st.image(r["img"], use_container_width=True)
                    st.caption("📷 **Quelle: SUVA**")
                else:
                    st.info("🖼️ Bild fehlt")
            with c2:
                st.subheader(r["titel"])
                st.write(r["desc"])
                st.caption("📝 **Basierend auf den lebenswichtigen Regeln der SUVA**")
    
    st.markdown("---")
    st.markdown("**Hinweis:** Diese App nutzt die lebenswichtigen Regeln der SUVA für interne Schulungs- und Sicherheitszwecke. Alle Materialien sind urheberrechtlich geschützt und Eigentum der SUVA. © SUVA")

elif st.session_state.current_page == 'bauav':
    st.header("⚖️ Bauarbeitenverordnung (BauAV)")
    st.markdown("**Nachschlagewerk für die wichtigsten Artikel der Schweizer Bauarbeitenverordnung**")
    
    # Suchfunktion
    st.markdown("---")
    search_query = st.text_input("🔍 Suche in BauAV", placeholder="z.B. Gerüst, Absturz, Leiter, Gräben...", help="Suchen Sie nach Artikelnummern, Titeln oder Begriffen im Text")
    st.markdown("---")
    
    def bauav_item(nr, titel, text, category=""):
        # Prüfe ob Artikel zur Suche passt
        if search_query:
            search_lower = search_query.lower()
            if (search_lower not in str(nr).lower() and 
                search_lower not in titel.lower() and 
                search_lower not in text.lower()):
                return False
        
        with st.expander(f"Art. {nr} - {titel}"):
            st.write(text)
        return True
    
    # Alle Artikel als Liste definieren
    artikel_liste = [
        # KATEGORIE 1: ORGANISATION & GRUNDLAGEN
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 3, "titel": "Planung und Organisation", "text": "Bauarbeiten müssen so geplant werden, dass das Risiko von Unfällen und Gesundheitsbeeinträchtigungen möglichst klein ist. Die Baustelle muss geordnet sein."},
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 4, "titel": "Kontrolle der Arbeitsmittel", "text": "Gerüste, Maschinen und Geräte müssen vor jedem Gebrauch auf Mängel geprüft werden. Defektes Material darf nicht verwendet werden."},
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 5, "titel": "Persönliche Schutzausrüstung (PSA)", "text": "Helmpflicht ist obligatorisch. Je nach Gefährdung sind Warnwesten, Sicherheitsschuhe, Gehörschutz oder Schutzbrillen zu tragen."},
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 6, "titel": "Verhalten bei Gefahr", "text": "Bei unmittelbarer Gefahr (z.B. drohender Einsturz, Unwetter) sind die Arbeiten sofort einzustellen und die Gefahrenzone zu verlassen."},
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 12, "titel": "Absperrung der Baustelle", "text": "Die Baustelle muss gegen unbefugtes Betreten gesichert sein. Zäune, Signale und Warnschilder sind erforderlich."},
        {"cat": 1, "cat_name": "Organisation & Grundlagen", "nr": 22, "titel": "Ordnung auf der Baustelle", "text": "Materialien sind stabil zu lagern. Keine Gefährdung durch Umkippen oder Wegrollen. Arbeitsplätze müssen aufgeräumt sein."},
        
        # KATEGORIE 2: ABSTURZSICHERUNG
        {"cat": 2, "cat_name": "Absturzsicherung & Öffnungen", "nr": 17, "titel": "Absturzkanten (Allgemein)", "text": "Ab einer Absturzhöhe von 2.00 m ist ein Seitenschutz zwingend (Holm, Zwischenholm, Bordbrett). Die Höhe des Seitenschutzes muss mindestens 1.00 m betragen."},
        {"cat": 2, "cat_name": "Absturzsicherung & Öffnungen", "nr": 25, "titel": "Bodenöffnungen", "text": "Löcher in Böden und Decken müssen durchbruchsicher abgedeckt und gegen Verschieben gesichert sein. Öffnungen sind deutlich zu kennzeichnen."},
        {"cat": 2, "cat_name": "Absturzsicherung & Öffnungen", "nr": 41, "titel": "Arbeiten an Dächern", "text": "Ab 2.00 m Absturzhöhe müssen Dächer durch Fassadengerüste, Spenglerläufe oder Auffangnetze gesichert werden. Steildächer ab 30° Neigung zusätzlich mit Seilsicherung."},
        {"cat": 2, "cat_name": "Absturzsicherung & Öffnungen", "nr": 19, "titel": "Herabfallende Gegenstände", "text": "Arbeitsbereiche, über denen gearbeitet wird, müssen gesichert sein (Schutzdächer oder Absperrungen). Werkzeuge müssen gegen Herunterfallen gesichert werden."},
        {"cat": 2, "cat_name": "Absturzsicherung & Öffnungen", "nr": 18, "titel": "Schutz der Personen unterhalb", "text": "Wenn Arbeiten in Höhe ausgeführt werden, muss der Bereich darunter abgesperrt oder mit Schutzdächern gesichert sein."},
        
        # KATEGORIE 3: ZUGÄNGE & LEITERN
        {"cat": 3, "cat_name": "Zugänge, Verkehrswege & Leitern", "nr": 10, "titel": "Verkehrswege", "text": "Wege müssen frei von Hindernissen sein. Stolperstellen (Kabel, Material) sind zu entfernen. Wege müssen ausreichend breit und beleuchtet sein."},
        {"cat": 3, "cat_name": "Zugänge, Verkehrswege & Leitern", "nr": 15, "titel": "Zugänge zu Arbeitsplätzen", "text": "Zugänge müssen sicher sein. Treppentürme sind Leitern vorzuziehen. Steigungen dürfen nicht zu steil sein (max. 45°)."},
        {"cat": 3, "cat_name": "Zugänge, Verkehrswege & Leitern", "nr": 21, "titel": "Verwendung von Leitern", "text": "Leitern dürfen nur für kurzzeitige Arbeiten verwendet werden. Sie sind gegen Wegrutschen zu sichern. Niemals auf der obersten Sprosse stehen."},
        {"cat": 3, "cat_name": "Zugänge, Verkehrswege & Leitern", "nr": 34, "titel": "Leitern (Bauart)", "text": "Anlegeleitern müssen die Austrittsstelle um mindestens 1.00 m überragen. Der Neigungswinkel sollte zwischen 65° und 75° liegen."},
        {"cat": 3, "cat_name": "Zugänge, Verkehrswege & Leitern", "nr": 14, "titel": "Treppen und Rampen", "text": "Treppen müssen mindestens 0.80 m breit sein und Handläufe aufweisen. Rampen dürfen nicht steiler als 15° sein."},
        
        # KATEGORIE 4: GERÜSTE
        {"cat": 4, "cat_name": "Gerüste", "nr": 47, "titel": "Gerüste (Allgemein)", "text": "Gerüste müssen standfest sein. Der Belag muss lückenlos verlegt sein. Änderungen dürfen nur vom Gerüstbauer vorgenommen werden. Tägliche Sichtkontrolle ist erforderlich."},
        {"cat": 4, "cat_name": "Gerüste", "nr": 57, "titel": "Rollgerüste", "text": "Rollgerüste dürfen nicht verschoben werden, solange sich Personen darauf befinden. Die Räder müssen arretiert sein. Maximale Höhe: 12 m."},
        {"cat": 4, "cat_name": "Gerüste", "nr": 48, "titel": "Gerüstbeläge", "text": "Beläge müssen durchbruchsicher sein. Überlappungen müssen mindestens 20 cm betragen. Keine schadhaften Bretter verwenden."},
        {"cat": 4, "cat_name": "Gerüste", "nr": 49, "titel": "Gerüstverankerung", "text": "Fassadengerüste müssen ausreichend verankert sein. Abstände der Verankerungen: alle 4 m in der Höhe, alle 6 m in der Breite."},
        {"cat": 4, "cat_name": "Gerüste", "nr": 50, "titel": "Gerüstmontage", "text": "Gerüste dürfen nur von qualifiziertem Personal errichtet werden. Standsicherheitsnachweis ist erforderlich bei Gerüsten über 3 m Höhe."},
        
        # KATEGORIE 5: TIEFBAU & GRÄBEN
        {"cat": 5, "cat_name": "Tiefbau & Gräben", "nr": 20, "titel": "Gräben und Schächte", "text": "Ab einer Tiefe von 1.50 m müssen Grabenwände gespriesst oder geböscht werden. Bei fliessenden Böden schon früher. Verbau muss durchbruchsicher sein."},
        {"cat": 5, "cat_name": "Tiefbau & Gräben", "nr": 82, "titel": "Arbeiten in der Nähe von Leitungen", "text": "Bei Grabarbeiten ist auf Werkleitungen (Gas, Strom, Wasser) zu achten. Pläne konsultieren! Mindestabstände beachten (Strom: 3-5 m je nach Spannung)."},
        {"cat": 5, "cat_name": "Tiefbau & Gräben", "nr": 23, "titel": "Erdarbeiten", "text": "Böschungen müssen stabil sein. Neigung maximal 45° bei bindigen Böden, 35° bei nichtbindigen Böden. Maschinenabstände von Grabenkanten beachten (min. 0.5 m)."},
        {"cat": 5, "cat_name": "Tiefbau & Gräben", "nr": 81, "titel": "Sprengarbeiten", "text": "Sprengarbeiten dürfen nur von qualifiziertem Personal ausgeführt werden. Sicherheitszone muss abgesperrt werden. Mindestabstand: 300 m."},
        
        # KATEGORIE 6: GESUNDHEIT & SPEZIELLES
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 32, "titel": "Schutz vor Sonne und Hitze", "text": "Arbeitsplätze sind wenn möglich zu beschatten. Den Mitarbeitern ist genügend Trinkwasser zur Verfügung zu stellen. Pausen an kühlen Orten einplanen."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 33, "titel": "Staub, Lärm, Vibrationen", "text": "Gesundheitsgefährdende Einwirkungen sind zu minimieren (z.B. Wasser gegen Staub, Gehörschutz bei Lärm über 85 dB(A)). Vibrationen durch Dämpfung reduzieren."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 83, "titel": "Elektrische Freileitungen", "text": "Für Baumaschinen gelten Mindestabstände zu Freileitungen (Niederspannung 3 m / Hochspannung 5 m+). Näherung nur mit speziellen Massnahmen."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 24, "titel": "Brandverhütung", "text": "Brennbare Materialien sicher lagern. Feuerlöscher müssen an gut zugänglichen Stellen bereitstehen. Rauchverbot auf Baustellen beachten."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 26, "titel": "Kranarbeiten", "text": "Krane müssen auf standsicherem Untergrund stehen. Ausleger nicht über Personen schwenken. Lasten sicher anschlagen (mindestens 4-fache Sicherheit)."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 27, "titel": "Hebearbeiten", "text": "Lasten nur von instruiertem Personal anschlagen. Niemals unter schwebenden Lasten stehen. Signale und Kommunikation zwischen Kranführer und Einweiser."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 28, "titel": "Schweissarbeiten", "text": "Schweissplätze müssen brandgeschützt eingerichtet sein. Brandwachen sind erforderlich. Sauerstoff und brennbare Gase getrennt lagern."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 29, "titel": "Umgang mit Chemikalien", "text": "Gefahrstoffe nach Sicherheitsdatenblatt handhaben. PSA entsprechend Gefährdung tragen. Behältnisse klar kennzeichnen."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 30, "titel": "Arbeitsplätze unter der Erde", "text": "Ausreichende Beleuchtung sicherstellen. Belüftung muss gewährleistet sein. Notausgänge kennzeichnen. Gasmessungen durchführen."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 31, "titel": "Lagerung von Materialien", "text": "Materialien stabil stapeln. Maximale Stapelhöhe beachten. Gänge zwischen Stapeln freihalten (min. 0.8 m)."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 35, "titel": "Baumaschinen", "text": "Maschinen nur von qualifiziertem Personal bedienen. Tägliche Sichtkontrolle erforderlich. Warntöne und Rückspiegel funktionsfähig halten."},
        {"cat": 6, "cat_name": "Gesundheit & Spezielle Gefahren", "nr": 36, "titel": "Fahrzeuge auf der Baustelle", "text": "Höchstgeschwindigkeit 10 km/h. Fußgängerbereiche kennzeichnen. Fahrzeuge müssen Tageslichtleuchten haben."},
    ]
    
    # Filtere Artikel basierend auf Suche
    if search_query:
        search_lower = search_query.lower()
        filtered_artikel = [
            art for art in artikel_liste
            if (search_lower in str(art["nr"]).lower() or
                search_lower in art["titel"].lower() or
                search_lower in art["text"].lower())
        ]
        if filtered_artikel:
            st.success(f"✅ {len(filtered_artikel)} Artikel gefunden für '{search_query}'")
        else:
            st.warning(f"⚠️ Keine Artikel gefunden für '{search_query}'. Versuchen Sie einen anderen Suchbegriff.")
    else:
        filtered_artikel = artikel_liste
    
    # Zeige Artikel nach Kategorien gruppiert
    current_cat = None
    displayed_count = 0
    
    for artikel in filtered_artikel:
        if artikel["cat"] != current_cat:
            if current_cat is not None:
                st.divider()
            st.markdown(f"### {artikel['cat']}. {artikel['cat_name']}")
            current_cat = artikel["cat"]
        
        with st.expander(f"Art. {artikel['nr']} - {artikel['titel']}"):
            st.write(artikel['text'])
        displayed_count += 1
    
    if search_query and displayed_count == 0:
        st.info("💡 **Tipp:** Suchen Sie nach Begriffen wie 'Gerüst', 'Absturz', 'Leiter', 'Gräben', 'PSA', etc.")

elif st.session_state.current_page == 'notfall':
    st.header("🚨 Notfallmanagement (SOS)")
    st.markdown("**Wenn etwas passiert, zählt jede Sekunde.**")
    st.markdown("---")
    
    st.subheader("📞 Notfallnummern")
    st.markdown("**Wählen Sie die richtige Nummer für Ihren Notfall:**")
    st.markdown("")
    
    # Notfallnummern in Spalten anzeigen mit Containern
    col1, col2 = st.columns(2)
    
    with col1:
        with st.container(border=True):
            st.markdown("### 🚑 144 - Sanitätsnotruf")
            st.markdown("**Wichtigste Nummer. Bei allen medizinischen Notfällen:**")
            st.markdown("- Unfall")
            st.markdown("- Herzinfarkt")
            st.markdown("- Sturz")
            st.markdown(f"[📞 144 anrufen](tel:144)", unsafe_allow_html=True)
        
        st.markdown("")
        
        with st.container(border=True):
            st.markdown("### 🚁 1414 - Rega (Luftrettung)")
            st.markdown("**Essenziell in der Schweiz. Bei:**")
            st.markdown("- Schwer zugänglichem Gelände")
            st.markdown("- Kran-Unfällen")
            st.markdown("- Wenn Bodenambulanzen zu lange brauchen")
            st.caption("ℹ️ *Hinweis: Im Wallis wird über die 144 disponiert, aber die 1414 ist national bekannt.*")
            st.markdown(f"[📞 1414 anrufen](tel:1414)", unsafe_allow_html=True)
        
        st.markdown("")
        
        with st.container(border=True):
            st.markdown("### 🚒 118 - Feuerwehr")
            st.markdown("**Nicht nur bei Feuer! Auch bei:**")
            st.markdown("- Personenrettung (aus Tiefen/Höhen)")
            st.markdown("- Chemieunfällen (Öl/Gefahrgut)")
            st.markdown("- Verschüttungen")
            st.markdown(f"[📞 118 anrufen](tel:118)", unsafe_allow_html=True)
    
    with col2:
        with st.container(border=True):
            st.markdown("### ☠️ 145 - Tox Info Suisse")
            st.markdown("**Bei Vergiftungen oder Unfällen mit Chemikalien/Baustoffen:**")
            st.markdown("- Verschlucken")
            st.markdown("- Einatmen")
            st.markdown("- Augenkontakt")
            st.markdown(f"[📞 145 anrufen](tel:145)", unsafe_allow_html=True)
        
        st.markdown("")
        
        with st.container(border=True):
            st.markdown("### 👮 117 - Polizei")
            st.markdown("**Bei:**")
            st.markdown("- Verkehrsunfällen vor der Baustelle")
            st.markdown("- Einbruch")
            st.markdown("- Gewaltandrohung")
            st.markdown(f"[📞 117 anrufen](tel:117)", unsafe_allow_html=True)
        
        st.markdown("")
        
        with st.container(border=True):
            st.markdown("### 🌍 112 - Euro-Notruf")
            st.markdown("**Funktioniert oft auch dann, wenn das eigene Handynetz kein Signal hat**")
            st.caption("*(Roaming über Fremdnetze)*")
            st.markdown(f"[📞 112 anrufen](tel:112)", unsafe_allow_html=True)
    
    st.markdown("---")
    st.subheader("❓ Die \"W-Fragen\"-Hilfe")
    st.info("💡 **Viele Leute stehen unter Schock. Ein kurzes Skript auf dem Bildschirm hilft:**")
    
    # W-Fragen in einem Streamlit Container
    with st.container(border=True):
        st.markdown("#### Beantworten Sie diese Fragen am Telefon:")
        st.markdown("")
        st.markdown("**Wer ruft an?**")
        st.caption("Ihr Name und Ihre Funktion")
        st.markdown("")
        st.markdown("**Wo ist es passiert?**")
        st.caption("Genauer Standort, Adresse, Baustelle")
        st.markdown("")
        st.markdown("**Was ist passiert?**")
        st.caption("Art des Unfalls, Verletzungen")
        st.markdown("")
        st.markdown("**Wie viele Verletzte?**")
        st.caption("Anzahl der betroffenen Personen")
    
    st.markdown("---")
    st.warning("⚠️ **Wichtig:** Bleiben Sie ruhig, sprechen Sie langsam und deutlich. Legen Sie nicht auf, bis die Rettungsleitstelle alle Informationen hat.")

elif st.session_state.current_page == 'gefahrstoff':
    st.header("🧪 Gefahrstoffkataster")
    st.markdown("**Digitaler Zugriff auf Sicherheitsdatenblätter für verwendete Chemikalien oder Baustoffe.**")
    
    # Wichtige Hinweise für die Praxis
    with st.expander("ℹ️ Wichtige Hinweise für die Praxis in der Schweiz", expanded=False):
        st.markdown("""
        **Sicherheitsdatenblatt-Pflicht:**
        - Sie müssen das Sicherheitsdatenblatt (SDB) nicht zwingend ausdrucken, aber jeder Mitarbeiter muss jederzeit digital Zugriff darauf haben (z.B. Tablet auf der Baustelle oder Ordner im Baucontainer).
        
        **Mengen-Schwelle:**
        - Es gibt keine "Mindestmenge", unter der man nichts dokumentieren muss. Aber: Für haushaltsübliche Produkte in Kleinstmengen (z.B. eine Tube Sekundenkleber oder Spülmittel) müssen Sie in der Regel kein Kataster führen. Sobald es gewerblich genutzt wird (der 10-Liter-Kanister Reiniger), gehört es hinein.
        
        **Substitution (Ersatzpflicht):**
        - Gemäss Schweizer Unfallversicherungsgesetz (UVG) müssen Sie prüfen: Kann ich diesen gefährlichen Stoff durch einen harmloseren ersetzen? (z.B. wasserbasierter Lack statt lösungsmittelhaltiger Lack). Wenn ja, müssen Sie das tun. Dokumentieren Sie diese Entscheidung kurz.
        """)
    
    st.markdown("---")
    
    gefahrstoffe = load_gefahrstoffe()
    
    # Suchfunktion
    search_query = st.text_input("🔍 Suche nach Gefahrstoff", placeholder="z.B. Beton, Lösungsmittel, Kleber...", help="Suchen Sie nach Name, Kategorie oder CAS-Nummer")
    st.markdown("---")
    
    # Tab-Layout
    tab1, tab2 = st.tabs(["📋 Gefahrstoffliste", "➕ Neuen Gefahrstoff hinzufügen"])
    
    with tab1:
        col_header1, col_header2 = st.columns([3, 1])
        with col_header1:
            st.subheader("Alle Gefahrstoffe")
        with col_header2:
            if is_admin():
                if st.button("🔄 Standard-Gefahrstoffe laden", use_container_width=True, help="Lädt die 6 gängigsten Gefahrstoffe aus dem Muster-Kataster"):
                    standard_gefahrstoffe = init_standard_gefahrstoffe()
                    # Nur hinzufügen, wenn noch nicht vorhanden (anhand Name prüfen)
                    existing_names = [g.get('name', '') for g in gefahrstoffe.values()]
                    added_count = 0
                    for std_id, std_data in standard_gefahrstoffe.items():
                        if std_data.get('name') not in existing_names:
                            gefahrstoffe[str(uuid.uuid4())] = std_data
                            added_count += 1
                    if added_count > 0:
                        save_gefahrstoffe(gefahrstoffe)
                        st.success(f"✅ {added_count} Standard-Gefahrstoff(e) hinzugefügt!")
                        st.rerun()
                    else:
                        st.info("ℹ️ Alle Standard-Gefahrstoffe sind bereits vorhanden.")
        
        if not gefahrstoffe:
            st.info("Noch keine Gefahrstoffe vorhanden. Fügen Sie einen neuen Gefahrstoff hinzu.")
        else:
            displayed_count = 0
            for gefahrstoff_id, gefahrstoff_data in gefahrstoffe.items():
                # Suchfilter anwenden (erweitert)
                if search_query:
                    search_lower = search_query.lower()
                    name = gefahrstoff_data.get('name', '').lower()
                    handelsbezeichnung = gefahrstoff_data.get('handelsbezeichnung', '').lower()
                    kategorie = gefahrstoff_data.get('kategorie', '').lower()
                    cas = gefahrstoff_data.get('cas_nummer', '').lower()
                    beschreibung = gefahrstoff_data.get('beschreibung', '').lower()
                    hersteller = gefahrstoff_data.get('hersteller', '').lower()
                    lagerort = gefahrstoff_data.get('lagerort', '').lower()
                    gefahrenbeschreibung = gefahrstoff_data.get('gefahrenbeschreibung', '').lower()
                    
                    if (search_lower not in name and 
                        search_lower not in handelsbezeichnung and
                        search_lower not in kategorie and 
                        search_lower not in cas and 
                        search_lower not in beschreibung and
                        search_lower not in hersteller and
                        search_lower not in lagerort and
                        search_lower not in gefahrenbeschreibung):
                        continue
                
                displayed_count += 1
                with st.container(border=True):
                    # Header mit Name und Löschen-Button
                    col_header1, col_header2 = st.columns([4, 1])
                    with col_header1:
                        st.markdown(f"### {gefahrstoff_data.get('handelsbezeichnung', gefahrstoff_data.get('name', 'Unbekannt'))}")
                    with col_header2:
                        if is_admin():
                            if st.button("🗑️ Löschen", key=f"del_{gefahrstoff_id}", use_container_width=True):
                                del gefahrstoffe[gefahrstoff_id]
                                save_gefahrstoffe(gefahrstoffe)
                                st.success("✅ Gefahrstoff gelöscht!")
                                st.rerun()
                    
                    # Erweiterte Informationen in Expander
                    with st.expander("📋 Alle Details anzeigen", expanded=False):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("#### Administrative Daten")
                            if gefahrstoff_data.get('handelsbezeichnung'):
                                st.markdown(f"**Handelsbezeichnung:** {gefahrstoff_data.get('handelsbezeichnung')}")
                            if gefahrstoff_data.get('hersteller'):
                                st.markdown(f"**Hersteller/Lieferant:** {gefahrstoff_data.get('hersteller')}")
                            if gefahrstoff_data.get('kategorie'):
                                st.markdown(f"**Kategorie:** {gefahrstoff_data.get('kategorie')}")
                            if gefahrstoff_data.get('cas_nummer'):
                                st.markdown(f"**CAS-Nummer:** {gefahrstoff_data.get('cas_nummer')}")
                            if gefahrstoff_data.get('lagerort'):
                                st.markdown(f"**Lagerort:** {gefahrstoff_data.get('lagerort')}")
                            if gefahrstoff_data.get('menge'):
                                st.markdown(f"**Menge:** {gefahrstoff_data.get('menge')}")
                            if gefahrstoff_data.get('sdb_datum'):
                                st.markdown(f"**SDB Datum:** {gefahrstoff_data.get('sdb_datum')}")
                        
                        with col2:
                            st.markdown("#### Gefahrenkennzeichnung")
                            if gefahrstoff_data.get('ghs_symbole'):
                                st.markdown(f"**GHS-Symbole:** {gefahrstoff_data.get('ghs_symbole')}")
                            if gefahrstoff_data.get('gefahrenbeschreibung'):
                                st.markdown(f"**Gefahrenbeschreibung:**")
                                st.write(gefahrstoff_data.get('gefahrenbeschreibung'))
                            if gefahrstoff_data.get('schutzmassnahmen'):
                                st.markdown(f"**Schutzmassnahmen (PSA & Technik):**")
                                st.write(gefahrstoff_data.get('schutzmassnahmen'))
                        
                        st.markdown("---")
                        st.markdown("#### Betriebsanweisung & Unterweisung")
                        col3, col4 = st.columns(2)
                        with col3:
                            if gefahrstoff_data.get('verwendung'):
                                st.markdown(f"**Verwendung:** {gefahrstoff_data.get('verwendung')}")
                            if gefahrstoff_data.get('betriebsanweisung_vorhanden'):
                                status = "✅ Ja" if gefahrstoff_data.get('betriebsanweisung_vorhanden') == "Ja" else "❌ Nein"
                                st.markdown(f"**Betriebsanweisung vorhanden:** {status}")
                        with col4:
                            if gefahrstoff_data.get('substitution'):
                                st.markdown(f"**Substitution (Ersatzpflicht):**")
                                st.write(gefahrstoff_data.get('substitution'))
                        
                        st.markdown("---")
                        st.markdown("#### Sicherheitsdatenblatt")
                        sdb_link = gefahrstoff_data.get('sdb_link', '')
                        sdb_datei = gefahrstoff_data.get('sdb_datei', '')
                        if sdb_link:
                            st.markdown(f"📄 [Sicherheitsdatenblatt öffnen]({sdb_link})", unsafe_allow_html=True)
                        elif sdb_datei:
                            st.markdown(f"📄 Sicherheitsdatenblatt: {sdb_datei}")
                        else:
                            st.caption("⚠️ Kein Sicherheitsdatenblatt hinterlegt")
                    
                    # Kurzübersicht (immer sichtbar)
                    col_short1, col_short2 = st.columns(2)
                    with col_short1:
                        if gefahrstoff_data.get('kategorie'):
                            st.caption(f"📦 {gefahrstoff_data.get('kategorie')}")
                        if gefahrstoff_data.get('lagerort'):
                            st.caption(f"📍 {gefahrstoff_data.get('lagerort')}")
                    with col_short2:
                        if gefahrstoff_data.get('ghs_symbole'):
                            st.caption(f"⚠️ {gefahrstoff_data.get('ghs_symbole')}")
                        sdb_link = gefahrstoff_data.get('sdb_link', '')
                        if sdb_link:
                            st.markdown(f"📄 [SDB öffnen]({sdb_link})", unsafe_allow_html=True)
            
            if search_query and displayed_count == 0:
                st.info("💡 **Tipp:** Keine Gefahrstoffe gefunden. Versuchen Sie eine andere Suchanfrage.")
    
    with tab2:
        if not is_admin():
            st.warning("⚠️ Nur Administratoren können neue Gefahrstoffe hinzufügen.")
        else:
            st.subheader("Neuen Gefahrstoff hinzufügen")
            
            with st.form("neuer_gefahrstoff", clear_on_submit=True):
                st.markdown("#### Administrative Daten (Pflicht)")
                handelsbezeichnung = st.text_input("Genaue Handelsbezeichnung *", help="z.B. 'Holcim Optimo 4' statt nur 'Zement'")
                hersteller = st.text_input("Hersteller/Lieferant *", help="Name der Firma")
                kategorie = st.selectbox("Kategorie *", ["Zementhaltige Produkte", "Lösungsmittelhaltige Farben/Lacke/Kleber", "Epoxidharze (2-Komponenten)", "PU-Produkte (Isocyanate)", "Kraftstoffe & Schmiermittel", "Reinigungsmittel (Sauer)", "Chemikalie", "Baustoff", "Klebstoff", "Lack/Farbe", "Reinigungsmittel", "Sonstiges"])
                cas_nummer = st.text_input("CAS-Nummer (optional)", help="Eindeutige Identifikationsnummer")
                lagerort = st.text_input("Lagerort *", help="z.B. 'Baucontainer A, Giftschrank' oder 'Magazin Regal 3'")
                menge = st.text_input("Durchschnittliche Lagermenge *", help="z.B. 'ca. 200 kg'")
                sdb_datum = st.date_input("SDB Datum (Erstellungsdatum des Sicherheitsdatenblatts)", value=None, help="Wenn älter als 5 Jahre -> neu anfordern")
                
                st.markdown("---")
                st.markdown("#### Gefahrenkennzeichnung")
                ghs_symbole = st.text_input("GHS-Symbole (Kennzeichnung)", help="z.B. 'GHS05, GHS07' für Ätzend/Reizend")
                gefahrenbeschreibung = st.text_area("Gefahrenbeschreibung (Risiken)", help="z.B. 'Verursacht schwere Augenschäden. Hautreizungen. Staub reizt Atemwege.'")
                schutzmassnahmen = st.text_area("Wichtige Schutzmassnahmen (PSA & Technik)", help="z.B. 'Handschuhe (Nitril/Butyl), Schutzbrille, lange Kleidung. Bei Staubentwicklung: Maske FFP2.'")
                
                st.markdown("---")
                st.markdown("#### Betriebsanweisung & Unterweisung")
                verwendung = st.text_input("Verwendung", help="Kurze Beschreibung, z.B. 'Verkleben von Bodenplatten'")
                betriebsanweisung_vorhanden = st.selectbox("Betriebsanweisung vorhanden?", ["Ja", "Nein"], help="Für Gefahrstoffe (besonders CMR-Stoffe) benötigen Sie eine schriftliche Anweisung")
                substitution = st.text_area("Substitution (Ersatzpflicht)", help="Kann dieser gefährliche Stoff durch einen harmloseren ersetzt werden? Dokumentieren Sie diese Entscheidung.")
                
                st.markdown("---")
                st.markdown("#### Sicherheitsdatenblatt")
                sdb_link = st.text_input("Link zum Sicherheitsdatenblatt (optional)", help="URL zum PDF oder Online-SDB. Jeder Mitarbeiter muss jederzeit digital Zugriff darauf haben.")
                
                submitted = st.form_submit_button("Gefahrstoff hinzufügen", use_container_width=True, type="primary")
                
                if submitted:
                    if not handelsbezeichnung or not hersteller or not lagerort or not menge:
                        st.error("❌ Bitte füllen Sie alle Pflichtfelder aus!")
                    else:
                        new_id = str(uuid.uuid4())
                        gefahrstoffe[new_id] = {
                            "name": handelsbezeichnung,  # Für Rückwärtskompatibilität
                            "handelsbezeichnung": handelsbezeichnung,
                            "hersteller": hersteller,
                            "kategorie": kategorie,
                            "cas_nummer": cas_nummer if cas_nummer else "",
                            "lagerort": lagerort,
                            "menge": menge,
                            "sdb_datum": sdb_datum.strftime('%d.%m.%Y') if sdb_datum else "",
                            "ghs_symbole": ghs_symbole if ghs_symbole else "",
                            "gefahrenbeschreibung": gefahrenbeschreibung if gefahrenbeschreibung else "",
                            "schutzmassnahmen": schutzmassnahmen if schutzmassnahmen else "",
                            "verwendung": verwendung if verwendung else "",
                            "betriebsanweisung_vorhanden": betriebsanweisung_vorhanden,
                            "substitution": substitution if substitution else "",
                            "sdb_link": sdb_link if sdb_link else "",
                            "sdb_datei": "",
                            "erstellt_am": date.today().strftime('%d.%m.%Y')
                        }
                        save_gefahrstoffe(gefahrstoffe)
                        st.success(f"✅ Gefahrstoff '{handelsbezeichnung}' erfolgreich hinzugefügt!")
                        st.rerun()

elif st.session_state.current_page == 'wetter':
    st.header("🌤️ Wetter-Warnungen")
    st.markdown("**Direkte Schnittstelle zu MeteoSchweiz für Sturmwarnungen (Kranbetrieb einstellen) oder Hitzewarnungen (SUVA Hitzemassnahmen).**")
    st.markdown("---")
    
    # MeteoSchweiz Links
    with st.container(border=True):
        st.markdown("### 📡 MeteoSchweiz")
        st.caption("Offizieller Wetterdienst der Schweiz")
        st.markdown("")
        
        # Direkte Links mit st.link_button
        try:
            st.link_button("📊 Wettervorhersage (Karte)", "https://www.meteoschweiz.admin.ch/#tab=forecast-map", use_container_width=True, type="primary")
            st.link_button("🌪️ Gefahrenkarte (Wetterwarnungen)", "https://www.meteoschweiz.admin.ch/service-und-publikationen/applikationen/gefahren.html#tab=severe-weather-map&weather-tab=all", use_container_width=True)
        except:
            # Fallback für ältere Streamlit-Versionen
            st.markdown('[📊 **Wettervorhersage (Karte)**](https://www.meteoschweiz.admin.ch/#tab=forecast-map)', unsafe_allow_html=True)
            st.markdown("")
            st.markdown('[🌪️ **Gefahrenkarte (Wetterwarnungen)**](https://www.meteoschweiz.admin.ch/service-und-publikationen/applikationen/gefahren.html#tab=severe-weather-map&weather-tab=all)', unsafe_allow_html=True)
        
        st.markdown("---")
        st.info("💡 **Hinweis:** Für detaillierte, aktuelle Warnungen besuchen Sie bitte die MeteoSchweiz-Website direkt. Installieren Sie die MeteoSwiss App auf Ihrem Smartphone für Push-Benachrichtigungen bei Warnungen.")
    
    st.markdown("---")
    
    # Wichtige Warnungen für Baustellen
    st.subheader("⚠️ Wichtige Warnungen für Baustellen")
    
    col_warn1, col_warn2 = st.columns(2)
    
    with col_warn1:
        with st.container(border=True):
            st.markdown("### 🌪️ Sturmwarnungen")
            st.markdown("**Kranbetrieb einstellen bei:**")
            st.markdown("")
            st.markdown("**🔴 Windgeschwindigkeit > 50 km/h (Bft 7)**")
            st.markdown("**🔴 Böen > 70 km/h**")
            st.markdown("**🔴 Warnung vor Sturm oder Orkan**")
            st.markdown("")
            st.error("⚠️ **SOFORTMASSNAHME:** Kranbetrieb sofort einstellen! Lasten sichern, Kran in Windrichtung ausrichten.")
            st.markdown("")
            st.markdown("**Weitere Massnahmen:**")
            st.markdown("• Lose Materialien sichern")
            st.markdown("• Gerüste prüfen (Verankerung)")
            st.markdown("• Baustelle absperren bei Gefahr")
    
    with col_warn2:
        with st.container(border=True):
            st.markdown("### ☀️ Hitzewarnungen (SUVA)")
            st.markdown("**Massnahmen bei Hitze:**")
            st.markdown("")
            st.markdown("**🟡 Temperaturen > 30°C:** Erhöhte Vorsicht")
            st.markdown("**🟠 Temperaturen > 35°C:** Zusätzliche Pausen")
            st.markdown("**🔴 Hitzewelle:** Anpassung der Arbeitszeiten")
            st.markdown("")
            st.warning("⚠️ **SUVA-Regeln:** Ausreichend trinken, Schattenplätze schaffen, Arbeitszeiten anpassen.")
            st.markdown("")
            st.markdown("**Weitere Massnahmen:**")
            st.markdown("• Genügend Trinkwasser bereitstellen")
            st.markdown("• Schattenplätze einrichten")
            st.markdown("• Arbeitszeiten anpassen (früher beginnen)")
            st.markdown("• PSA anpassen (luftdurchlässige Kleidung)")
    
    st.markdown("---")
    
    # Praktische Checkliste
    st.subheader("📋 Checkliste: Wetter-Check vor Baustellenstart")
    
    with st.container(border=True):
        col_check1, col_check2 = st.columns(2)
        
        with col_check1:
            st.markdown("#### Vor Arbeitsbeginn prüfen:")
            st.markdown("- ☐ Aktuelle Wetterwarnungen abrufen")
            st.markdown("- ☐ Windgeschwindigkeit prüfen (Kranbetrieb?)")
            st.markdown("- ☐ Temperatur prüfen (Hitzemassnahmen?)")
            st.markdown("- ☐ Niederschlagswahrscheinlichkeit")
            st.markdown("- ☐ Gewitterwarnung vorhanden?")
        
        with col_check2:
            st.markdown("#### Bei Warnungen:")
            st.markdown("- ☐ Baustellenleiter informieren")
            st.markdown("- ☐ Massnahmen umsetzen (Kran stoppen, etc.)")
            st.markdown("- ☐ Mitarbeiter informieren")
            st.markdown("- ☐ PSA anpassen")
            st.markdown("- ☐ Arbeitszeiten anpassen")
    
    st.markdown("---")
    
    # SUVA Hitzemassnahmen
    st.subheader("🌡️ SUVA Hitzemassnahmen (Detail)")
    
    with st.expander("📖 Detaillierte SUVA-Richtlinien für Hitze", expanded=False):
        st.markdown("""
        **Bei Temperaturen über 30°C:**
        - Regelmässige Pausen im Schatten (alle 1-2 Stunden)
        - Mindestens 0.5 Liter Wasser pro Stunde trinken
        - Leichte, luftdurchlässige Kleidung tragen
        - Kopfbedeckung verwenden
        
        **Bei Temperaturen über 35°C:**
        - Arbeitszeiten anpassen (früher beginnen, Mittagspause verlängern)
        - Schwere körperliche Arbeiten vermeiden
        - Zusätzliche Pausen (alle 30-60 Minuten)
        - Überwachung der Mitarbeiter (Anzeichen von Hitzschlag)
        
        **Symptome eines Hitzschlags:**
        - Kopfschmerzen, Schwindel, Übelkeit
        - Rote, heisse, trockene Haut
        - Verwirrtheit, Bewusstlosigkeit
        
        **Erste Hilfe bei Hitzschlag:**
        - Sofort in den Schatten bringen
        - Kühlen (feuchte Tücher, Wasser)
        - Notruf 144 wählen
        """)
    
    st.markdown("---")
    st.info("💡 **Wichtig:** Diese Seite dient als Schnittstelle zu offiziellen Wetterdiensten. Für aktuelle, verbindliche Warnungen konsultieren Sie bitte immer die offiziellen Quellen (MeteoSchweiz oder search.ch).")

elif st.session_state.current_page == 'kunden':
    if not is_admin():
        st.error("⛔ Zugriff verweigert. Diese Seite ist nur für Administratoren verfügbar.")
        st.info("Bitte als Admin einloggen, um auf die Kundenverwaltung zuzugreifen.")
    else:
        st.header("👥 Kundenverwaltung")
        st.markdown("---")
        
        customers = load_customers()
        
        # Tab-Layout
        tab1, tab2 = st.tabs(["📋 Kundenliste", "➕ Neuen Kunden hinzufügen"])
        
        with tab1:
            st.subheader("Alle Kunden")
            if not customers:
                st.info("Noch keine Kunden vorhanden. Fügen Sie einen neuen Kunden hinzu.")
            else:
                users = load_users()
                for kunde_id, kunde_data in customers.items():
                    with st.container(border=True):
                        col1, col2, col3 = st.columns([3, 1, 1])
                        with col1:
                            st.markdown(f"### {kunde_data.get('name', 'Unbekannt')}")
                            st.write(f"**Firma:** {kunde_data.get('firma', '-')}")
                            email = kunde_data.get('email', '')
                            username = kunde_data.get('username', '')
                            st.write(f"**Email:** {email}")
                            if username:
                                st.write(f"**Benutzername:** {username}")
                            st.write(f"**Telefon:** {kunde_data.get('telefon', '-')}")
                            if 'adresse' in kunde_data:
                                st.write(f"**Adresse:** {kunde_data['adresse']}")
                            
                            # Credits anzeigen
                            credits = int(kunde_data.get('credits', 0))
                            st.metric("🪙 SafeSite Credits", credits)
                            
                            # Login-Status anzeigen
                            has_login_email = email and email in users
                            has_login_username = username and username in users
                            if has_login_email or has_login_username:
                                st.success("✅ Login aktiv")
                                login_info = []
                                if has_login_email:
                                    login_info.append(f"Email: {email}")
                                if has_login_username:
                                    login_info.append(f"Username: {username}")
                                st.caption(" | ".join(login_info))
                            else:
                                st.warning("⚠️ Kein Login erstellt")
                        with col2:
                            has_login_email = email and email in users
                            has_login_username = username and username in users
                            if has_login_email or has_login_username:
                                if st.button("🔑 Passwort ändern", key=f"passwd_{kunde_id}"):
                                    st.session_state[f"edit_passwd_{kunde_id}"] = True
                                    st.rerun()
                            else:
                                if st.button("🔑 Login erstellen", key=f"create_login_{kunde_id}"):
                                    st.session_state[f"create_login_{kunde_id}"] = True
                                    st.rerun()
                            # Credits bearbeiten Button
                            if st.button("💰 Credits verwalten", key=f"credits_{kunde_id}"):
                                st.session_state[f"edit_credits_{kunde_id}"] = True
                                st.rerun()
                        with col3:
                            if st.button("🗑️ Löschen", key=f"delete_{kunde_id}"):
                                # Kunde aus customers.json löschen
                                del customers[kunde_id]
                                save_customers(customers)
                                # Login aus users.json löschen (falls vorhanden)
                                if email and email in users:
                                    del users[email]
                                if username and username in users:
                                    del users[username]
                                save_users(users)
                                st.success("Kunde gelöscht!")
                                st.rerun()
                        
                        # Credits bearbeiten Formular
                        if st.session_state.get(f"edit_credits_{kunde_id}", False):
                            st.divider()
                            with st.form(f"form_credits_{kunde_id}"):
                                st.markdown("**🪙 Credits verwalten**")
                                current_credits = int(kunde_data.get('credits', 0))
                                new_credits = st.number_input("Anzahl Credits", min_value=0, value=current_credits, step=1, key=f"credits_input_{kunde_id}")
                                col_a, col_b = st.columns(2)
                                with col_a:
                                    if st.form_submit_button("✅ Credits speichern", use_container_width=True):
                                        update_customer_credits(kunde_id, new_credits)
                                        st.session_state[f"edit_credits_{kunde_id}"] = False
                                        st.success(f"Credits auf {new_credits} aktualisiert!")
                                        st.rerun()
                                with col_b:
                                    if st.form_submit_button("❌ Abbrechen", use_container_width=True):
                                        st.session_state[f"edit_credits_{kunde_id}"] = False
                                        st.rerun()
                        
                        # Passwort ändern Formular
                        if st.session_state.get(f"edit_passwd_{kunde_id}", False):
                            st.divider()
                            with st.form(f"form_passwd_{kunde_id}"):
                                new_pass = st.text_input("Neues Passwort", type="password", key=f"new_pass_{kunde_id}")
                                new_pass_confirm = st.text_input("Passwort bestätigen", type="password", key=f"new_pass_confirm_{kunde_id}")
                                col_a, col_b = st.columns(2)
                                with col_a:
                                    if st.form_submit_button("✅ Passwort ändern", use_container_width=True):
                                        if new_pass and new_pass == new_pass_confirm:
                                            # Aktualisiere Passwort für Email-Login (falls vorhanden)
                                            if email and email in users:
                                                users[email] = new_pass
                                            # Aktualisiere Passwort für Username-Login (falls vorhanden)
                                            if username and username in users:
                                                users[username] = new_pass
                                            save_users(users)
                                            st.session_state[f"edit_passwd_{kunde_id}"] = False
                                            st.success("Passwort erfolgreich geändert!")
                                            st.rerun()
                                        elif new_pass != new_pass_confirm:
                                            st.error("Passwörter stimmen nicht überein!")
                                        else:
                                            st.error("Passwort darf nicht leer sein!")
                                with col_b:
                                    if st.form_submit_button("❌ Abbrechen", use_container_width=True):
                                        st.session_state[f"edit_passwd_{kunde_id}"] = False
                                        st.rerun()
                        
                        # Login erstellen Formular
                        if st.session_state.get(f"create_login_{kunde_id}", False):
                            st.divider()
                            with st.form(f"form_create_login_{kunde_id}"):
                                st.info(f"Login wird für: {kunde_data.get('name', 'Kunde')} erstellt")
                                login_options = []
                                if email:
                                    login_options.append(f"Email: {email}")
                                if username:
                                    login_options.append(f"Benutzername: {username}")
                                if login_options:
                                    st.caption(" | ".join(login_options))
                                new_pass = st.text_input("Passwort", type="password", key=f"create_pass_{kunde_id}")
                                new_pass_confirm = st.text_input("Passwort bestätigen", type="password", key=f"create_pass_confirm_{kunde_id}")
                                col_a, col_b = st.columns(2)
                                with col_a:
                                    if st.form_submit_button("✅ Login erstellen", use_container_width=True):
                                        if new_pass and new_pass == new_pass_confirm:
                                            # Erstelle Login mit Email (immer)
                                            if email:
                                                users[email] = new_pass
                                            # Erstelle auch Login mit Benutzername (falls vorhanden)
                                            if username:
                                                users[username] = new_pass
                                            save_users(users)
                                            st.session_state[f"create_login_{kunde_id}"] = False
                                            login_info = []
                                            if email:
                                                login_info.append(f"Email: {email}")
                                            if username:
                                                login_info.append(f"Username: {username}")
                                            st.success(f"Login erfolgreich erstellt! ({' | '.join(login_info)})")
                                            st.rerun()
                                        elif new_pass != new_pass_confirm:
                                            st.error("Passwörter stimmen nicht überein!")
                                        else:
                                            st.error("Passwort darf nicht leer sein!")
                                with col_b:
                                    if st.form_submit_button("❌ Abbrechen", use_container_width=True):
                                        st.session_state[f"create_login_{kunde_id}"] = False
                                        st.rerun()
                        
                        st.divider()
        
        with tab2:
            st.subheader("Neuen Kunden hinzufügen")
            with st.form("neuer_kunde", clear_on_submit=True):
                kunde_name = st.text_input("Name *", placeholder="Max Mustermann")
                firma = st.text_input("Firma", placeholder="Mustermann AG")
                email = st.text_input("Email *", placeholder="max@mustermann.ch")
                username_optional = st.text_input("Benutzername (optional)", placeholder="max.mustermann", help="Optional: Falls leer, kann sich der Kunde nur mit Email einloggen. Falls gesetzt, kann er sich mit Email ODER Benutzername einloggen.")
                telefon = st.text_input("Telefon", placeholder="+41 79 123 45 67")
                adresse = st.text_area("Adresse", placeholder="Musterstrasse 123\n8000 Zürich")
                
                st.divider()
                st.markdown("**🪙 SafeSite Credits:**")
                initial_credits = st.number_input("Anfangliche Credits", min_value=0, value=0, step=1, key="new_kunde_credits")
                st.caption("💡 1 Credit = 1 Bericht. Credits werden automatisch bei jedem Bericht abgebucht.")
                
                st.divider()
                st.markdown("**🔐 Login für SafeSite-Check (Pflicht):**")
                st.info("💡 Ein Login-Konto wird automatisch für jeden Kunden erstellt.")
                login_passwort = st.text_input("Passwort *", type="password", key="new_kunde_pass", help="Pflichtfeld: Der Kunde benötigt ein Passwort für den Login")
                login_passwort_confirm = st.text_input("Passwort bestätigen *", type="password", key="new_kunde_pass_confirm")
                login_info = []
                if email:
                    login_info.append(f"Email: {email}")
                if username_optional:
                    login_info.append(f"Benutzername: {username_optional}")
                if login_info:
                    st.caption(f"💡 Der Kunde kann sich mit {' oder '.join(login_info)} anmelden.")
                else:
                    st.caption("💡 Bitte Email (und optional Benutzername) eingeben.")
                
                col1, col2 = st.columns(2)
                with col1:
                    submit = st.form_submit_button("✅ Kunde hinzufügen", use_container_width=True)
                with col2:
                    cancel = st.form_submit_button("❌ Abbrechen", use_container_width=True)
                
                if submit:
                    if not kunde_name or not email:
                        st.error("⚠️ Name und Email sind Pflichtfelder!")
                    elif not login_passwort:
                        st.error("⚠️ Passwort ist ein Pflichtfeld!")
                    elif login_passwort != login_passwort_confirm:
                        st.error("⚠️ Die Passwörter stimmen nicht überein!")
                    else:
                        # Prüfen ob Email oder Benutzername bereits als Login existiert
                        users = load_users()
                        if email in users:
                            st.error(f"⚠️ Ein Login mit der Email '{email}' existiert bereits!")
                        elif username_optional and username_optional in users:
                            st.error(f"⚠️ Ein Login mit dem Benutzernamen '{username_optional}' existiert bereits!")
                        else:
                            # Eindeutige ID generieren
                            kunde_id = str(uuid.uuid4())[:8]
                            
                            # Kunde hinzufügen
                            customer_data = {
                                "name": kunde_name,
                                "firma": firma,
                                "email": email,
                                "telefon": telefon,
                                "adresse": adresse,
                                "credits": int(initial_credits),
                                "erstellt_am": date.today().strftime('%d.%m.%Y')
                            }
                            # Benutzername nur hinzufügen, wenn angegeben
                            if username_optional:
                                customer_data["username"] = username_optional
                            customers[kunde_id] = customer_data
                            save_customers(customers)
                            
                            # Login automatisch erstellen (immer, da Pflicht)
                            # Login mit Email (immer)
                            users[email] = login_passwort
                            # Login mit Benutzername (falls vorhanden)
                            if username_optional:
                                users[username_optional] = login_passwort
                            save_users(users)
                            
                            login_info = [f"Email: {email}"]
                            if username_optional:
                                login_info.append(f"Benutzername: {username_optional}")
                            st.success(f"✅ Kunde '{kunde_name}' erfolgreich hinzugefügt mit {initial_credits} Credits und Login erstellt ({' | '.join(login_info)})!")
                            st.rerun()
