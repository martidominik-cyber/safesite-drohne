"""
SafeSite Drohne – Berichterstellung (PDF & Word)
Format: Exakt wie SSD_Bericht Vorlage
  - Logo im Header (rechts)
  - Titel: SICHERHEITS-INSPEKTION (DROHNE)
  - Metadaten: Projekt, Datum, Inspektor, Status
  - Orange Überschrift: ZUSAMMENFASSUNG / MÄNGEL (#F47E0E)
  - Pro Mangel: Orange Nummer + Kategorie, dann Mangel/Verstoss/Massnahme, Bild
  - Freigabe-Sektion mit Unterschriftslinien
"""
import os
import time
import cv2
from datetime import date
from config import LOGO_FILE

# SafeSite Orange
ORANGE_HEX = "F47E0E"
ORANGE_RGB = (244, 126, 14)

# ============================================================
# HILFSFUNKTIONEN
# ============================================================
def extract_frame(video_path: str, timestamp: float):
    try:
        cap = cv2.VideoCapture(video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        total = cap.get(cv2.CAP_PROP_FRAME_COUNT)
        duration = total / fps if fps > 0 else 30
        timestamp = max(0, min(float(timestamp), duration - 0.5))
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
        ret, frame = cap.read()
        cap.release()
        if ret:
            return cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    except Exception:
        pass
    return None


def _safe(text) -> str:
    """Entfernt Zeichen, die fpdf nicht darstellen kann."""
    if text is None:
        return ""
    return str(text).encode("latin-1", "ignore").decode("latin-1")


def _get_image(item, index, m_type, m_files):
    """Gibt (img_path, is_temp) zurück."""
    if m_type == "video":
        frame = extract_frame(m_files[0], item.get("zeitstempel_sekunden", 0))
        if frame is not None:
            path = f"temp_report_{index}.jpg"
            cv2.imwrite(path, cv2.cvtColor(frame, cv2.COLOR_RGB2BGR))
            return path, True
    elif m_type == "images":
        idx = item.get("bild_index", 0)
        if idx < len(m_files):
            return m_files[idx], False
    return None, False


# ============================================================
# WORD-BERICHT (python-docx)
# ============================================================
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


def create_word(data, m_type, m_files, projekt, inspektor, status):
    """Erstellt Word-Bericht im SSD-Format."""
    if not WORD_AVAILABLE:
        return None

    doc = Document()

    # --- Seitenränder ---
    for section in doc.sections:
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # --- Logo im Header (rechts) ---
    if os.path.exists(LOGO_FILE):
        try:
            header = doc.sections[0].header
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = hp.add_run()
            run.add_picture(LOGO_FILE, width=Cm(4))
        except Exception:
            pass

    # --- Titel ---
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("SICHERHEITS-INSPEKTION (DROHNE)")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # --- Metadaten ---
    meta_p = doc.add_paragraph()
    _add_meta_line(meta_p, "Projekt:", f" {projekt}")
    meta_p.add_run("\n")
    _add_meta_line(meta_p, "Datum:", f" {date.today().strftime('%d.%m.%Y')}")
    meta_p.add_run("\n")
    _add_meta_line(meta_p, "Inspektor:", f" {inspektor}")
    meta_p.add_run("\n")
    _add_meta_line(meta_p, "Status:", f" {status}")

    # --- Überschrift: ZUSAMMENFASSUNG / MÄNGEL ---
    doc.add_paragraph()  # Leerzeile
    summary_p = doc.add_paragraph()
    summary_run = summary_p.add_run("ZUSAMMENFASSUNG / MÄNGEL")
    summary_run.font.size = Pt(14)
    summary_run.font.color.rgb = RGBColor(*ORANGE_RGB)
    summary_run.underline = True

    # --- Mängel ---
    temps = []
    for i, item in enumerate(data):
        # Kategorie-Überschrift (orange)
        cat_p = doc.add_paragraph()
        cat_run = cat_p.add_run(f"{i+1}. {item.get('kategorie', 'Mangel')}")
        cat_run.font.size = Pt(14)
        cat_run.font.color.rgb = RGBColor(*ORANGE_RGB)

        # Verification-Label (falls vorhanden)
        verif_label = item.get("verifikation_label", "")
        if verif_label:
            verif_p = doc.add_paragraph()
            verif_run = verif_p.add_run(verif_label)
            verif_run.font.size = Pt(9)
            verif_run.font.color.rgb = RGBColor(100, 100, 100)

        # Mangel / Verstoss / Massnahme
        content_p = doc.add_paragraph()

        # Mangel
        _add_meta_line(content_p, "Mangel:", f" {item.get('mangel', '-')}")
        content_p.add_run("\n\n")

        # Verstoss
        _add_meta_line(content_p, "Verstoss:", f" {item.get('verstoss', '-')}")
        content_p.add_run("\n\n")

        # Massnahme
        _add_meta_line(content_p, "Massnahme:", f" {item.get('massnahme', '-')}")

        # Bild
        img_path, is_temp = _get_image(item, i, m_type, m_files)
        if img_path:
            try:
                doc.add_picture(img_path, width=Inches(5.5))
            except Exception:
                pass
            if is_temp:
                temps.append(img_path)

        doc.add_paragraph()  # Leerzeile

    # --- FREIGABE ---
    doc.add_paragraph()
    freigabe_nr = len(data) + 1
    freigabe_p = doc.add_paragraph()
    freigabe_run = freigabe_p.add_run(f"{freigabe_nr}. FREIGABE")
    freigabe_run.font.size = Pt(14)
    freigabe_run.font.color.rgb = RGBColor(*ORANGE_RGB)

    doc.add_paragraph("Dieser Bericht wurde generiert durch SafeSite Drohne.")

    # Hinweis (kursiv)
    hint_p = doc.add_paragraph()
    hint_run = hint_p.add_run(
        "Hinweis: Dieser Bericht dient als visuelle Unterstützung. "
        "Er entbindet die zuständigen nicht von der gesetzlichen Kontrollpflicht."
    )
    hint_run.italic = True
    hint_run.font.size = Pt(9)

    doc.add_paragraph()

    # Unterschriftslinien
    sig1 = doc.add_paragraph()
    sig1.add_run(f"Erstellt durch: {inspektor} ").bold = False
    sig1.add_run("____________________")

    doc.add_paragraph()
    sig2 = doc.add_paragraph()
    sig2.add_run("Verantwortlicher: ").bold = False
    sig2.add_run("____________________")

    out = "Bericht.docx"
    doc.save(out)

    for t in temps:
        try:
            os.remove(t)
        except Exception:
            pass
    return out


def _add_meta_line(paragraph, label: str, value: str):
    """Fügt bold Label + normal Value zu einem Paragraph hinzu."""
    run_label = paragraph.add_run(label)
    run_label.bold = True
    paragraph.add_run(value)


# ============================================================
# PDF-BERICHT (fpdf2)
# ============================================================
from fpdf import FPDF


class _PDF(FPDF):
    def header(self):
        if os.path.exists(LOGO_FILE):
            try:
                self.image(LOGO_FILE, 155, 8, 40)
            except Exception:
                pass
        self.ln(25)


def create_pdf(data, m_type, m_files, projekt, inspektor, status):
    """Erstellt PDF-Bericht im SSD-Format."""
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # --- Titel ---
    pdf.set_font("Arial", "B", 22)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 12, "SICHERHEITS-INSPEKTION (DROHNE)", ln=True)
    pdf.ln(6)

    # --- Metadaten ---
    _pdf_meta(pdf, "Projekt:", _safe(projekt))
    _pdf_meta(pdf, "Datum:", date.today().strftime("%d.%m.%Y"))
    _pdf_meta(pdf, "Inspektor:", _safe(inspektor))

    # Status mit Farbe
    pdf.set_font("Arial", "B", 11)
    pdf.cell(30, 7, "Status:", ln=0)
    pdf.set_font("Arial", "", 11)
    if "Massnahmen" in status:
        pdf.set_text_color(255, 153, 51)
    elif "Kritisch" in status:
        pdf.set_text_color(204, 0, 0)
    else:
        pdf.set_text_color(0, 153, 0)
    pdf.cell(0, 7, _safe(status), ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    # --- Überschrift: ZUSAMMENFASSUNG / MÄNGEL ---
    pdf.set_font("Arial", "BU", 14)
    pdf.set_text_color(*ORANGE_RGB)
    pdf.cell(0, 10, "ZUSAMMENFASSUNG / MAENGEL", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # --- Mängel ---
    temps = []
    for i, item in enumerate(data):
        if pdf.get_y() > 230:
            pdf.add_page()

        # Kategorie-Überschrift (orange)
        pdf.set_font("Arial", "B", 14)
        pdf.set_text_color(*ORANGE_RGB)
        cat = f"{i+1}. {_safe(item.get('kategorie', 'Mangel'))}"
        pdf.cell(0, 10, cat, ln=True)
        pdf.set_text_color(0, 0, 0)

        # Verification-Label
        verif_label = item.get("verifikation_label", "")
        if verif_label:
            pdf.set_font("Arial", "I", 8)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, _safe(verif_label), ln=True)
            pdf.set_text_color(0, 0, 0)

        # Mangel
        pdf.set_font("Arial", "B", 10)
        pdf.cell(22, 6, "Mangel:", ln=0)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, _safe(item.get("mangel", "-")))
        pdf.ln(3)

        # Verstoss
        if pdf.get_y() > 250:
            pdf.add_page()
        pdf.set_font("Arial", "B", 10)
        pdf.cell(22, 6, "Verstoss:", ln=0)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, _safe(item.get("verstoss", "-")))
        pdf.ln(3)

        # Massnahme
        if pdf.get_y() > 250:
            pdf.add_page()
        pdf.set_font("Arial", "B", 10)
        pdf.cell(28, 6, "Massnahme:", ln=0)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, _safe(item.get("massnahme", "-")))
        pdf.ln(4)

        # Bild
        img_path, is_temp = _get_image(item, i, m_type, m_files)
        if img_path:
            if pdf.get_y() > 160:
                pdf.add_page()
            try:
                pdf.image(img_path, x=20, w=140)
            except Exception:
                pass
            pdf.ln(8)
            if is_temp:
                temps.append(img_path)

        pdf.ln(4)

    # --- FREIGABE ---
    if pdf.get_y() > 200:
        pdf.add_page()
    pdf.ln(10)

    freigabe_nr = len(data) + 1
    pdf.set_font("Arial", "B", 14)
    pdf.set_text_color(*ORANGE_RGB)
    pdf.cell(0, 10, f"{freigabe_nr}. FREIGABE", ln=True)
    pdf.set_text_color(0, 0, 0)

    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 8, "Dieser Bericht wurde generiert durch SafeSite Drohne.", ln=True)

    pdf.set_font("Arial", "I", 9)
    pdf.multi_cell(0, 5,
        "Hinweis: Dieser Bericht dient als visuelle Unterstuetzung. "
        "Er entbindet die zustaendigen nicht von der gesetzlichen Kontrollpflicht.")
    pdf.ln(15)

    # Unterschriftslinien
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 10, f"Erstellt durch: {_safe(inspektor)} ____________________", ln=True)
    pdf.ln(5)
    pdf.cell(0, 10, "Verantwortlicher: ____________________", ln=True)

    out = "Bericht.pdf"
    pdf.output(out)

    for t in temps:
        try:
            os.remove(t)
        except Exception:
            pass
    return out


def _pdf_meta(pdf, label: str, value: str):
    """Schreibt eine Metadaten-Zeile im PDF."""
    pdf.set_font("Arial", "B", 11)
    pdf.cell(30, 7, label, ln=0)
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 7, value, ln=True)
